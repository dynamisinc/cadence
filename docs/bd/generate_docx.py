"""Generate CAPABILITIES.docx from the narrative capabilities document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(18)
style_h1.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
style_h1.font.bold = True

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(14)
style_h2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
style_h2.font.bold = True

style_h3 = doc.styles['Heading 3']
style_h3.font.size = Pt(12)
style_h3.font.color.rgb = RGBColor(0x2D, 0x5A, 0x8A)
style_h3.font.bold = True


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A3C6E')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'EDF2F9')

    return table


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CADENCE')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Platform Capabilities')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2D, 0x5A, 0x8A)

doc.add_paragraph('')

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('HSEEP-Compliant Exercise Design, Conduct, and Evaluation Platform')
run.font.size = Pt(13)
run.italic = True
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Cadence is a purpose-built web platform for the design, conduct, and evaluation of emergency '
    'management exercises in full compliance with the Homeland Security Exercise and Evaluation '
    'Program (HSEEP) 2020 doctrine. The platform provides end-to-end management of Master Scenario '
    'Events Lists (MSELs), real-time exercise conduct with dual-time tracking, structured HSEEP '
    'evaluation through Exercise Evaluation Guides (EEGs), and comprehensive after-action review support.'
)

doc.add_paragraph(
    'Cadence was designed from the ground up to address the operational realities of exercise conduct: '
    'unreliable field connectivity, multi-hour sessions that cannot tolerate data loss, multi-agency '
    'participation requiring strict data isolation, and the need for real-time coordination across '
    'geographically distributed teams. The platform serves Exercise Directors, Controllers, Evaluators, '
    'Observers, Players, Simulators, Facilitators, Safety Officers, and Trusted Agents \u2014 every '
    'HSEEP-defined participant role \u2014 with role-appropriate interfaces and permissions.'
)

doc.add_paragraph(
    'The platform operates as a multi-tenant, cloud-hosted SaaS application with offline-first '
    'architecture, ensuring uninterrupted exercise conduct regardless of network conditions.'
)

# ============================================================
# 2. HSEEP COMPLIANCE
# ============================================================
doc.add_heading('2. HSEEP Compliance', level=1)

doc.add_heading('2.1 Doctrinal Alignment', level=2)
doc.add_paragraph(
    'Cadence implements the HSEEP 2020 exercise lifecycle with native support for all HSEEP-defined '
    'exercise types, participant roles, evaluation methodology, and documentation requirements.'
)

doc.add_heading('Supported Exercise Types', level=3)
add_table(doc,
    ['Type', 'Abbreviation', 'Description'],
    [
        ['Tabletop Exercise', 'TTX', 'Discussion-based exercises with Facilitator-Paced timing mode'],
        ['Functional Exercise', 'FE', 'Operations-based exercises with simulated response'],
        ['Full-Scale Exercise', 'FSE', 'Operations-based exercises deploying real resources'],
        ['Computer-Aided Exercise', 'CAX', 'Exercises integrated with simulation systems'],
    ]
)

doc.add_paragraph('')
doc.add_heading('HSEEP Participant Roles', level=3)
add_table(doc,
    ['Role', 'Platform Capabilities'],
    [
        ['Exercise Director', 'Full exercise authority: configuration, activation, participant management, Go/No-Go decisions, AAR review'],
        ['Controller', 'MSEL delivery, inject firing, scenario flow management, field observation capture'],
        ['Evaluator', 'Structured EEG assessment, ad-hoc observation capture, P/S/M/U rating, objective-linked evaluation'],
        ['Observer', 'Read-only exercise monitoring without interference'],
        ['Player', 'Inject receipt and acknowledgment in virtual/hybrid exercises'],
        ['Simulator', 'Simulated communications and SimCell activity management'],
        ['Facilitator', 'Discussion flow guidance and pace control for TTX exercises'],
        ['Safety Officer', 'Safety oversight with exercise pause/stop authority'],
        ['Trusted Agent', 'Embedded subject matter expertise with observation capabilities'],
    ]
)

doc.add_paragraph('')
doc.add_heading('2.2 HSEEP Inject Status Workflow', level=2)
doc.add_paragraph(
    'Cadence implements the FEMA PrepToolkit inject status workflow as its default lifecycle:'
)
add_para(doc, 'Draft \u2192 Submitted \u2192 Approved \u2192 Synchronized \u2192 Released \u2192 Complete', bold=True)
doc.add_paragraph(
    'With parallel paths for Deferred (cancelled before delivery) and Obsolete (soft-removed but '
    'retained for audit). The platform supports a formal inject approval workflow with configurable '
    'policies at the organization, exercise, and inject levels, ensuring quality control and governance '
    'compliance before exercise conduct.'
)

doc.add_heading('2.3 Exercise Evaluation Guide (EEG) Implementation', level=2)
doc.add_paragraph(
    'Cadence implements the full HSEEP evaluation chain:'
)
add_para(doc, 'Objective \u2192 Capability \u2192 Capability Target \u2192 Critical Task \u2192 Inject \u2192 EEG Entry', bold=True)
doc.add_paragraph(
    'Exercise planners define Capability Targets with measurable performance thresholds and Critical '
    'Tasks that specify the actions required to achieve each target. MSEL injects are linked to the '
    'Critical Tasks they are designed to test. During conduct, Evaluators record structured EEG entries '
    'against specific Critical Tasks using the HSEEP P/S/M/U rating scale:'
)

add_table(doc,
    ['Rating', 'Definition'],
    [
        ['P', 'Performed without Challenges'],
        ['S', 'Performed with Some Challenges'],
        ['M', 'Performed with Major Challenges'],
        ['U', 'Unable to be Performed'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph(
    'An EEG Coverage Dashboard provides real-time visibility into which Critical Tasks have been evaluated, '
    'the performance distribution by Capability Target, and gaps requiring attention. All EEG data exports '
    'in AAR-ready format organized by Capability, Target, Task, and Observation.'
)

doc.add_heading('2.4 Cross-Domain Framework Support', level=2)
doc.add_paragraph(
    'While HSEEP is the default framework, Cadence provides configurable status workflows and '
    'terminology to serve organizations operating under alternative frameworks:'
)
add_table(doc,
    ['Framework', 'Audience'],
    [
        ['DoD / Joint Training System (JTS)', 'Military exercises (STARTEX/ENDEX, JMSEL)'],
        ['NATO', 'Allied military coordination (EXCON/DISTAFF, LIVEX/CPX)'],
        ['UK Cabinet Office', 'UK government (Main Events List, Cell structure)'],
        ['Australian AIIMS', 'Australia/NZ emergency management'],
        ['NIST / MITRE ATT&CK', 'Cybersecurity exercises (Red/Blue/Purple teams)'],
        ['CMS / Joint Commission', 'Healthcare exercises (HICS, surge levels)'],
        ['FFIEC / FINRA', 'Financial sector exercises (RTO/RPO/MTD metrics)'],
        ['ISO 22301 / BCI', 'Private sector business continuity'],
    ]
)

# ============================================================
# 3. EXERCISE DESIGN AND PLANNING
# ============================================================
doc.add_paragraph('')
doc.add_heading('3. Exercise Design and Planning', level=1)

doc.add_heading('3.1 Exercise Lifecycle Management', level=2)
doc.add_paragraph(
    'Cadence manages exercises through a complete HSEEP-compliant lifecycle:'
)
add_para(doc, 'Draft \u2192 Active \u2192 Paused \u2192 Completed \u2192 Archived', bold=True)
doc.add_paragraph(
    'Each status enforces appropriate data editability, access controls, and available actions. '
    'The platform supports pause and resume capabilities for administrative holds, safety events, '
    'and multi-day exercises. Exercise status and clock state operate independently, providing '
    'flexibility for scenarios such as administrative pauses while scenario time continues, or '
    'clock pauses for discussion while the exercise remains active.'
)
doc.add_paragraph(
    'Transition safeguards include validation rules (e.g., exercises cannot activate without at '
    'least one inject), confirmation dialogs for destructive transitions, and a full audit trail '
    'of who performed each transition and when.'
)

doc.add_heading('3.2 MSEL Authoring', level=2)
doc.add_paragraph(
    'The Master Scenario Events List is the core artifact of exercise planning. Cadence provides:'
)
p = doc.add_paragraph()
run = p.add_run('Inject Management: ')
run.bold = True
p.add_run(
    'Full create, read, update, and delete operations for injects with support for all standard '
    'MSEL fields including inject number (auto-sequential), title, description, scheduled time, '
    'scenario time, sender, recipient, delivery method, expected player action, and controller notes.'
)

p = doc.add_paragraph()
run = p.add_run('Dual-Time Tracking: ')
run.bold = True
p.add_run(
    'Every inject supports two independent time concepts \u2014 Scheduled Time (wall clock delivery) '
    'and Scenario Time (in-story time). This enables multi-day scenarios to be compressed into shorter '
    'exercise windows.'
)

p = doc.add_paragraph()
run = p.add_run('Phase Organization: ')
run.bold = True
p.add_run(
    'Exercises can be structured into named phases representing time periods, operational stages, '
    'or scenario segments. Injects are assigned to phases and the MSEL can be grouped and collapsed '
    'by phase for large exercises.'
)

p = doc.add_paragraph()
run = p.add_run('Objective Linkage: ')
run.bold = True
p.add_run(
    'HSEEP-compliant SMART objectives are defined per exercise and linked to injects in a many-to-many '
    'relationship, enabling coverage analysis to verify all objectives are adequately exercised.'
)

p = doc.add_paragraph()
run = p.add_run('Inject Filtering and Search: ')
run.bold = True
p.add_run(
    'Multi-criteria filtering by status, phase, objective, delivery method, and controller assignment. '
    'Full-text search across inject fields for rapid location of specific content in large MSELs.'
)

p = doc.add_paragraph()
run = p.add_run('Inject Organization: ')
run.bold = True
p.add_run(
    'Configurable sorting, grouping by phase/status/objective with collapsible sections, and '
    'drag-and-drop reordering that automatically updates sequence numbers.'
)

doc.add_heading('3.3 Excel Import and Export', level=2)
doc.add_paragraph(
    'Cadence preserves existing organizational workflows by providing full round-trip Excel compatibility:'
)
p = doc.add_paragraph()
run = p.add_run('Import: ')
run.bold = True
p.add_run(
    'A guided wizard supports uploading Excel (.xlsx, .xls) and CSV files, mapping spreadsheet columns '
    'to Cadence inject fields with intelligent synonym matching for delivery methods, validating data '
    'before finalization, and creating or updating injects by inject number.'
)

p = doc.add_paragraph()
run = p.add_run('Export: ')
run.bold = True
p.add_run(
    'MSEL data exports to formatted Excel workbooks with consistent column ordering that matches the '
    'import template for round-trip compatibility. Exports include optional metadata sheets with exercise '
    'information, objectives, and phases. Blank templates are available for download.'
)

doc.add_heading('3.4 Exercise Configuration', level=2)
doc.add_paragraph(
    'Before an exercise begins, Cadence guides setup through a structured configuration workflow '
    'with a visual progress dashboard covering: Basic Information, Participant Assignment, '
    'Time Zone Configuration, Clock Mode Selection, Objective Definition, Phase Definition, '
    'MSEL Population, and Inject Approval. The progress dashboard displays completion status for '
    'each configuration area and prevents exercise activation until required items are complete.'
)

doc.add_heading('3.5 Inject Approval Workflow', level=2)
doc.add_paragraph(
    'Organizations can require formal approval of injects before exercise conduct. The workflow is '
    'configurable at three tiers: Organization Level (Disabled, Optional, or Required), Exercise Level '
    '(Directors enable/disable per exercise), and Inject Level (status workflow enforced per setting). '
    'When enabled, injects follow a Draft \u2192 Submitted \u2192 Approved workflow with an Approval '
    'Queue view, batch approval capabilities, configurable approval permissions, self-approval policies, '
    'and a Go-Live Gate that prevents activation until all injects are approved. Every status change is '
    'recorded with full audit trail.'
)

doc.add_heading('3.6 MSEL Duplication and Reuse', level=2)
doc.add_paragraph(
    'Organizations running recurring exercises can duplicate existing exercises including their complete '
    'MSEL, preserving inject content, phase structure, and objective definitions while creating a new '
    'exercise for modification.'
)

# ============================================================
# 4. EXERCISE CONDUCT
# ============================================================
doc.add_heading('4. Exercise Conduct', level=1)

doc.add_heading('4.1 Real-Time Exercise Clock', level=2)
doc.add_paragraph(
    'Cadence provides two timing modes to support different exercise types:'
)
p = doc.add_paragraph()
run = p.add_run('Clock-Driven Mode: ')
run.bold = True
p.add_run(
    'For operations-based exercises (FE, FSE), the exercise clock runs in real time. Injects '
    'automatically transition to "Ready" status when the clock reaches their scheduled delivery time. '
    'Controllers see injects organized into Now, Upcoming, and Completed sections.'
)
p = doc.add_paragraph()
run = p.add_run('Facilitator-Paced Mode: ')
run.bold = True
p.add_run(
    'For discussion-based exercises (TTX), the Facilitator manually advances through injects in '
    'sequence without time constraints, controlling the pace of discussion.'
)

doc.add_heading('4.2 Inject Delivery (Firing)', level=2)
doc.add_paragraph(
    'Controllers deliver injects through the platform\'s firing mechanism. When a Controller fires '
    'an inject, the status transitions to Released, both wall clock and scenario time are recorded, '
    'all connected users receive real-time notification, evaluators are alerted to begin observation, '
    'and the exercise timeline updates across all clients. Fire confirmation dialogs are available '
    'for critical injects.'
)

doc.add_heading('4.3 Real-Time Synchronization', level=2)
doc.add_paragraph(
    'Cadence uses Azure SignalR Service for sub-second synchronization across all connected users. '
    'Inject status changes, clock events, observations, and participant activity propagate to all '
    'clients instantly.'
)

doc.add_heading('4.4 Offline Capability and Field Resilience', level=2)
doc.add_paragraph(
    'Cadence was architected from day one for environments with unreliable connectivity, directly '
    'addressing the operational reality that exercises frequently occur in Emergency Operations Centers, '
    'field locations, and disaster sites with poor network access.'
)
p = doc.add_paragraph()
run = p.add_run('Offline Detection: ')
run.bold = True
p.add_run('Continuous connectivity monitoring with clear visual indicators of connection status.')

p = doc.add_paragraph()
run = p.add_run('Local Data Cache: ')
run.bold = True
p.add_run('All active exercise data cached locally in IndexedDB for full read access regardless of connectivity.')

p = doc.add_paragraph()
run = p.add_run('Offline Action Queue: ')
run.bold = True
p.add_run(
    'User actions queued in a persistent FIFO queue stored in IndexedDB with optimistic UI updates.'
)

p = doc.add_paragraph()
run = p.add_run('Sync on Reconnect: ')
run.bold = True
p.add_run(
    'Queued actions processed sequentially with conflict detection. Last-write-wins for most operations; '
    'first-write-wins for inject firing.'
)

p = doc.add_paragraph()
run = p.add_run('Zero Data Loss Guarantee: ')
run.bold = True
p.add_run(
    'All queued operations survive browser crashes and are reconciled on reconnect with user notification '
    'of any conflicts.'
)

doc.add_heading('4.5 Notifications', level=2)
doc.add_paragraph(
    'A real-time notification system with toast alerts for immediate events, a persistent notification '
    'bell with unread count, priority-based display, role-targeted delivery, and database-backed '
    'persistence across sessions.'
)

doc.add_heading('4.6 Field Operations', level=2)
doc.add_paragraph(
    'For operations-based exercises where participants are deployed in the field, Cadence provides '
    'photo capture and attachment with client-side compression, voice-to-text observation input, '
    'opt-in GPS location tracking with Director map view and automatic geo-stamping, and safety-flagged '
    'observations with immediate Director and Safety Officer visibility.'
)

# ============================================================
# 5. EVALUATION AND OBSERVATION
# ============================================================
doc.add_heading('5. Evaluation and Observation', level=1)

doc.add_heading('5.1 Ad-Hoc Observation Capture', level=2)
doc.add_paragraph(
    'Evaluators capture real-time observations through a quick-entry interface with type classification '
    '(Strength, Area for Improvement, Neutral), optional P/S/M/U rating, inject and objective linkage, '
    'dual timestamps, and offline support.'
)

doc.add_heading('5.2 Structured EEG Assessment', level=2)
doc.add_paragraph(
    'Evaluators perform structured assessments against the Exercise Evaluation Guide by selecting '
    'a Capability Target and Critical Task, recording observation text, assigning a mandatory P/S/M/U '
    'rating, and optionally linking to the triggering inject.'
)

doc.add_heading('5.3 Evaluation Coverage Dashboard', level=2)
doc.add_paragraph(
    'Real-time visibility during and after conduct showing overall Critical Task coverage, P/S/M/U '
    'rating distribution by Capability Target, and identification of unevaluated tasks for Director '
    'action.'
)

# ============================================================
# 6. AFTER-ACTION REVIEW AND REPORTING
# ============================================================
doc.add_heading('6. After-Action Review and Reporting', level=1)

doc.add_heading('6.1 Review Mode', level=2)
doc.add_paragraph(
    'A dedicated post-conduct view with phase-grouped timeline, inject outcome summaries with '
    'time variance analysis, observation review panel, and exercise statistics dashboard optimized '
    'for AAR discussion.'
)

doc.add_heading('6.2 Metrics and Analytics', level=2)
p = doc.add_paragraph()
run = p.add_run('Exercise-Level Metrics: ')
run.bold = True
p.add_run(
    'Real-time inject delivery statistics, on-time delivery rate, observation counts and P/S/M/U '
    'distribution, evaluator coverage, controller activity, and timeline analysis.'
)
p = doc.add_paragraph()
run = p.add_run('Organization-Level Metrics: ')
run.bold = True
p.add_run(
    'Cross-exercise performance trends, capability tracking, comparative analysis between exercises, '
    'benchmark comparison, and custom configurable dashboards. All metrics export to PDF, Excel, and '
    'image formats.'
)

doc.add_heading('6.3 AAR Export', level=2)
doc.add_paragraph(
    'EEG-based AAR export organizes findings in the HSEEP-required structure: '
    'Capability \u2192 Capability Target \u2192 Critical Task \u2192 Observations/EEG Entries. '
    'Each entry includes P/S/M/U rating, evaluator identity, timestamp, linked inject, and full '
    'observation text, directly supporting the HSEEP AAR/IP format.'
)

# ============================================================
# 7. MULTI-TENANCY AND ORGANIZATION MANAGEMENT
# ============================================================
doc.add_heading('7. Multi-Tenancy and Organization Management', level=1)

doc.add_paragraph(
    'Cadence implements a multi-tenant architecture with Organization as the primary security and data '
    'isolation boundary. All exercise data, user assignments, configurations, and audit trails are '
    'scoped to an organization. Data from one organization is never visible to another.'
)

doc.add_heading('Three-Tier Role Hierarchy', level=3)
add_table(doc,
    ['Tier', 'Scope', 'Roles'],
    [
        ['System', 'Platform-wide', 'Admin, Manager, User'],
        ['Organization', 'Per-organization', 'OrgAdmin, OrgManager, OrgUser'],
        ['Exercise', 'Per-exercise', 'All 9 HSEEP roles'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph(
    'Users can belong to multiple organizations with different roles in each, supporting consultants '
    'and contractors who work across agencies. Organization features include lifecycle management, '
    'user invitation via email or shareable codes, agency list management, capability library selection '
    '(FEMA, NATO, NIST, ISO, custom), configurable defaults, and complete audit trails.'
)

# ============================================================
# 8. AUTHENTICATION AND SECURITY
# ============================================================
doc.add_heading('8. Authentication and Security', level=1)

doc.add_paragraph(
    'Cadence implements a hybrid authentication architecture supporting local credentials and '
    'Azure Entra SSO with automatic account linking. The token strategy uses 15-minute access '
    'tokens stored in memory (preventing XSS attacks) and 4-hour refresh tokens in HttpOnly cookies. '
    'Session duration is aligned with typical exercise conduct. Authorization is enforced server-side '
    'on every request with exercise-scoped permissions and role inheritance.'
)

# ============================================================
# 9. PLATFORM EXPERIENCE
# ============================================================
doc.add_heading('9. Platform Experience', level=1)

doc.add_paragraph(
    'Cadence provides a role-aware, context-adaptive user interface with a personalized dashboard, '
    'context-aware navigation that transforms when entering an exercise, persistent exercise clock '
    'display, responsive design for desktop and tablet, extended session support (4+ hours), '
    'auto-save, comprehensive keyboard shortcuts, and a three-tier settings model (user, exercise, '
    'organization) with auto-save and reset-to-default at every level.'
)

# ============================================================
# 10. TECHNICAL ARCHITECTURE
# ============================================================
doc.add_heading('10. Technical Architecture', level=1)

add_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend', 'React 19, TypeScript 5, Material UI 7', 'Responsive single-page application'],
        ['Build System', 'Vite 7', 'Fast development and production builds'],
        ['Backend API', '.NET 10, ASP.NET Core', 'RESTful API (always-warm, no cold starts)'],
        ['ORM', 'Entity Framework Core 10', 'Database access with migrations'],
        ['Database', 'Azure SQL', 'Relational data storage'],
        ['Real-Time', 'Azure SignalR Service', 'Sub-second synchronization'],
        ['Background Jobs', 'Azure Functions', 'Timer-triggered maintenance tasks'],
        ['Offline Storage', 'IndexedDB (Dexie.js)', 'Client-side data cache and action queue'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph(
    'The REST API runs on Azure App Service (not serverless) to eliminate cold starts during exercise '
    'conduct. Azure Functions are used only for background timer tasks. The entire offline architecture '
    'was designed before the first feature was built, ensuring every capability works offline by default. '
    'Deployment uses CI/CD via GitHub Actions with infrastructure defined as Bicep templates.'
)

# ============================================================
# 11. DIFFERENTIATING CAPABILITIES
# ============================================================
doc.add_heading('11. Differentiating Capabilities', level=1)

add_table(doc,
    ['Capability', 'Description'],
    [
        ['Offline-First Architecture', 'Full functionality during connectivity loss with zero data loss guarantee and automatic sync on reconnect.'],
        ['Dual-Time Tracking', 'Native support for both wall clock and scenario time on every inject, enabling multi-day scenario compression.'],
        ['HSEEP EEG Implementation', 'Structured evaluation chain from Objectives through Capability Targets, Critical Tasks, and EEG Entries with P/S/M/U ratings.'],
        ['Cross-Domain Framework Support', 'Configurable workflows for HSEEP, DoD/JTS, NATO, cybersecurity, healthcare, financial, and international frameworks.'],
        ['Independent Clock and Status', 'Exercise status and clock state operate independently, supporting complex operational patterns.'],
        ['Nine HSEEP Roles', 'Complete implementation of all HSEEP-defined participant roles with exercise-scoped permissions.'],
        ['Three-Tier Inject Approval', 'Configurable approval workflows at organization, exercise, and inject levels with batch operations and Go-Live gate.'],
        ['Field Operations Suite', 'Photo capture, voice-to-text, GPS location tracking, and real-time Director situational awareness.'],
        ['Multi-Tenant Data Isolation', 'Organization-level security boundary with validated data isolation across tenants.'],
        ['Always-Warm API', 'No cold starts during exercise conduct. The API is always responsive for time-sensitive inject delivery.'],
    ]
)

# Save
output_path = r'C:\Code\dynamis\cadence\docs\bd\Cadence_Platform_Capabilities.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
