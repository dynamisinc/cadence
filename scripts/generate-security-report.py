#!/usr/bin/env python3
"""
Generate Cadence Application Security Assessment Report as a Word document.
Parses consolidated SARIF from the automated security pipeline and produces
a professional report suitable for customers, compliance evaluators, and auditors.
"""

import json
import os
from datetime import datetime
from collections import Counter, defaultdict

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# =============================================================================
# Configuration
# =============================================================================

# Paths — resolve relative to repo root (script expects to run from repo root)
SARIF_PATH = os.environ.get("SARIF_PATH", "/tmp/security-report/consolidated-security-report.sarif")
LOGO_PATH = "src/frontend/public/icons/icon-192x192.png"
OUTPUT_PATH = "docs/Cadence-Security-Assessment-Report.docx"
SCAN_DATE = os.environ.get("SCAN_DATE", datetime.now().strftime("%Y-%m-%d"))

# Brand colors
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xD3, 0x2F, 0x2F)
ORANGE = RGBColor(0xF5, 0x7C, 0x00)
BLUE = RGBColor(0x19, 0x76, 0xD2)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
TEAL = RGBColor(0x00, 0x79, 0x6B)

SEVERITY_COLORS = {"error": RED, "warning": ORANGE, "note": BLUE, "none": MED_GRAY}
# Tool-reported severity labels (raw SARIF level) — used as secondary metadata only
TOOL_SEVERITY_LABELS = {"error": "Error", "warning": "Warning", "note": "Note", "none": "Info"}
# Assessed risk labels — used as the PRIMARY classification in the report
RISK_COLORS = {"Critical": RED, "High": RED, "Medium": ORANGE, "Low": BLUE, "Informational": MED_GRAY, "Mitigated": GREEN}
RISK_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4, "Mitigated": 5}
# Legacy alias kept for backward compat
SEVERITY_LABELS = TOOL_SEVERITY_LABELS

# CWE to OWASP Top 10 2025 mapping
CWE_TO_OWASP = {
    "CWE-319": ("A04:2025", "Cryptographic Failures"),
    "CWE-524": ("A02:2025", "Security Misconfiguration"),
    "CWE-352": ("A02:2025", "Security Misconfiguration"),
    "CWE-16": ("A02:2025", "Security Misconfiguration"),
    "CWE-693": ("A02:2025", "Security Misconfiguration"),
}

# Checkov rule to OWASP/CWE mapping (OWASP Top 10 2025)
CHECKOV_OWASP = {
    "CKV_AZURE_67": ("A02:2025", "Security Misconfiguration", "CWE-16"),
    "CKV_AZURE_16": ("A07:2025", "Authentication Failures", "CWE-306"),
    "CKV_AZURE_212": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_18": ("A02:2025", "Security Misconfiguration", "CWE-16"),
    "CKV_AZURE_213": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_71": ("A07:2025", "Authentication Failures", "CWE-306"),
    "CKV_AZURE_17": ("A07:2025", "Authentication Failures", "CWE-295"),
    "CKV_AZURE_222": ("A02:2025", "Security Misconfiguration", "CWE-284"),
    "CKV_AZURE_225": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV2_AZURE_27": ("A07:2025", "Authentication Failures", "CWE-306"),
    "CKV_AZURE_113": ("A02:2025", "Security Misconfiguration", "CWE-284"),
    "CKV_AZURE_229": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_27": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_26": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_35": ("A02:2025", "Security Misconfiguration", "CWE-284"),
    "CKV_AZURE_43": ("A02:2025", "Security Misconfiguration", "CWE-16"),
    "CKV_AZURE_206": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_84": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_19": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_87": ("A02:2025", "Security Misconfiguration", "CWE-693"),
    "CKV_AZURE_21": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_22": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_20": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_24": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_25": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
    "CKV_AZURE_23": ("A09:2025", "Logging & Alerting Failures", "CWE-778"),
}

# Checkov rule descriptions (human-readable)
CHECKOV_DESCRIPTIONS = {
    "CKV_AZURE_67": "The Azure Function App should be configured to use the latest HTTP version (HTTP/2) for improved performance and security features.",
    "CKV_AZURE_16": "The App Service should register with Microsoft Entra ID to enable centralized identity management and conditional access policies.",
    "CKV_AZURE_212": "The App Service should have a minimum of two instances to ensure high availability and failover capability during maintenance or outages.",
    "CKV_AZURE_18": "The web application should use the latest HTTP version (HTTP/2) to benefit from multiplexing, header compression, and improved security.",
    "CKV_AZURE_213": "The App Service should configure a health check endpoint to enable automatic instance replacement when an instance becomes unhealthy.",
    "CKV_AZURE_71": "The web application should use a managed identity (system-assigned or user-assigned) instead of credentials for authenticating to Azure services.",
    "CKV_AZURE_17": "Enabling client certificate authentication on the App Service adds an additional layer of mutual TLS authentication for incoming requests.",
    "CKV_AZURE_222": "Public network access to the App Service should be restricted. Access should be limited to a virtual network or specific IP ranges.",
    "CKV_AZURE_225": "The App Service Plan should be configured for zone redundancy to ensure availability across Azure availability zones.",
    "CKV2_AZURE_27": "Microsoft Entra ID authentication should be enabled for Azure SQL to replace SQL-based authentication with centralized identity management.",
    "CKV_AZURE_113": "The Azure SQL Server should disable public network access and use private endpoints for network isolation.",
    "CKV_AZURE_229": "The Azure SQL Database should be configured for zone redundancy to protect against datacenter-level failures.",
    "CKV_AZURE_27": "Email notifications for SQL Server security alerts should include service and co-administrators to ensure timely incident response.",
    "CKV_AZURE_26": "SQL Server security alert notifications should be configured with specific recipient email addresses for targeted alerting.",
    "CKV_AZURE_35": "The Storage Account default network access rule should be set to Deny, with explicit allow rules for trusted networks only.",
    "CKV_AZURE_43": "The Storage Account name should follow Azure naming conventions and organizational naming standards.",
    "CKV_AZURE_206": "The Storage Account should use geo-redundant replication (GRS or RA-GRS) to protect against regional outages.",
    "CKV_AZURE_84": "Azure Defender for Storage should be enabled to detect unusual and potentially harmful access or exploitation attempts.",
    "CKV_AZURE_19": "Azure Defender for Cloud should use the Standard pricing tier to enable advanced threat protection across resources.",
    "CKV_AZURE_87": "Azure Defender for Key Vault should be enabled to detect unusual access patterns and potential compromise of secrets.",
    "CKV_AZURE_21": "Email notifications for high-severity Defender for Cloud alerts should be enabled to ensure rapid incident awareness.",
    "CKV_AZURE_22": "High-severity alert email notifications should also be sent to subscription owners for visibility at the management level.",
    "CKV_AZURE_20": "A security contact phone number should be configured in Defender for Cloud for urgent security incident escalation.",
    "CKV_AZURE_24": "SQL Server auditing retention should be configured for at least 90 days to meet compliance requirements and enable forensic analysis.",
    "CKV_AZURE_25": "SQL Server threat detection should enable alerts for all threat types including SQL injection, brute force, and anomalous access.",
    "CKV_AZURE_23": "SQL Server auditing should be enabled to track database events and maintain an audit trail for compliance and security investigations.",
}

# Risk classification for each finding type
RISK_CLASSIFICATION = {
    # Checkov - grouped by risk
    "CKV_AZURE_222": "Medium",
    "CKV_AZURE_113": "Medium",
    "CKV_AZURE_35": "Mitigated",  # Storage networkAcls defaultAction: 'Deny' added
    "CKV_AZURE_23": "Mitigated",  # SQL auditing enabled in database.bicep
    "CKV_AZURE_24": "Mitigated",  # SQL audit retention set to 90 days
    "CKV_AZURE_25": "Mitigated",  # SQL threat detection enabled
    "CKV2_AZURE_27": "Medium",
    "CKV_AZURE_71": "Mitigated",  # SystemAssigned managed identity added to webapp
    "CKV_AZURE_16": "Low",
    "CKV_AZURE_17": "Low",
    "CKV_AZURE_18": "Low",
    "CKV_AZURE_67": "Low",
    "CKV_AZURE_213": "Low",
    "CKV_AZURE_84": "Mitigated",  # Defender for Storage Standard tier enabled
    "CKV_AZURE_87": "Mitigated",  # Defender for Key Vault Standard tier enabled
    "CKV_AZURE_19": "Mitigated",  # Defender for SQL Standard tier enabled
    "CKV_AZURE_27": "Low",
    "CKV_AZURE_26": "Low",
    "CKV_AZURE_21": "Low",
    "CKV_AZURE_22": "Low",
    "CKV_AZURE_20": "Low",
    "CKV_AZURE_43": "Informational",
    "CKV_AZURE_206": "Informational",
    "CKV_AZURE_212": "Informational",
    "CKV_AZURE_225": "Informational",
    "CKV_AZURE_229": "Informational",
    # ZAP
    "zap-10035": "Mitigated",  # HSTS header now configured in ASP.NET Core middleware
    "zap-10038": "Mitigated",  # CSP header now set on all API responses
    "zap-10049": "Informational",
    "zap-90005": "Informational",
}


# =============================================================================
# SARIF Parser
# =============================================================================

def parse_sarif(path):
    """Parse SARIF and return structured findings."""
    with open(path) as f:
        sarif = json.load(f)

    tools = []
    all_findings = []

    for run in sarif.get("runs", []):
        driver = run["tool"]["driver"]
        tool_name = driver["name"]
        tool_version = driver.get("semanticVersion", driver.get("version", "N/A"))
        rules_map = {}
        for rule in driver.get("rules", []):
            rules_map[rule["id"]] = rule

        tools.append({
            "name": tool_name,
            "version": tool_version,
            "rules_count": len(rules_map),
            "results_count": len(run.get("results", [])),
        })

        for result in run.get("results", []):
            rule_id = result.get("ruleId", "unknown")
            level = result.get("level", "none")
            message = result.get("message", {}).get("text", "")
            locations = []
            for loc in result.get("locations", []):
                phys = loc.get("physicalLocation", {})
                uri = phys.get("artifactLocation", {}).get("uri", "")
                region = phys.get("region", {})
                line = region.get("startLine", "")
                method = loc.get("properties", {}).get("method", "")
                locations.append({"uri": uri, "line": line, "method": method})

            rule_info = rules_map.get(rule_id, {})
            short_desc = rule_info.get("shortDescription", {}).get("text", message.split("|")[0].strip())
            full_desc = rule_info.get("fullDescription", {}).get("text", "")
            help_text = rule_info.get("help", {}).get("text", "")
            tags = rule_info.get("properties", {}).get("tags", [])
            confidence = rule_info.get("properties", {}).get("confidence", 0)

            all_findings.append({
                "tool": tool_name,
                "tool_version": tool_version,
                "rule_id": rule_id,
                "level": level,
                "message": message,
                "short_desc": short_desc,
                "full_desc": full_desc,
                "help_text": help_text,
                "tags": tags,
                "confidence": confidence,
                "locations": locations,
            })

    return tools, all_findings


def deduplicate_findings(findings):
    """Deduplicate findings by (tool, level, ruleId). Return list of unique findings with counts."""
    groups = defaultdict(lambda: {"count": 0, "locations": [], "finding": None})
    for f in findings:
        key = (f["tool"], f["level"], f["rule_id"])
        groups[key]["count"] += 1
        groups[key]["locations"].extend(f["locations"])
        if groups[key]["finding"] is None:
            groups[key]["finding"] = f

    deduped = []
    for key, group in groups.items():
        entry = dict(group["finding"])
        entry["instance_count"] = group["count"]
        # Deduplicate locations by URI
        seen_uris = set()
        unique_locs = []
        for loc in group["locations"]:
            loc_key = f"{loc['method']} {loc['uri']}:{loc['line']}"
            if loc_key not in seen_uris:
                seen_uris.add(loc_key)
                unique_locs.append(loc)
        entry["unique_locations"] = unique_locs
        deduped.append(entry)

    # Sort: error first, then warning, then note
    level_order = {"error": 0, "warning": 1, "note": 2, "none": 3}
    deduped.sort(key=lambda x: (level_order.get(x["level"], 9), x["rule_id"]))
    return deduped


# =============================================================================
# Document Helpers
# =============================================================================

def setup_document():
    """Create and configure the Word document."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.color.rgb = NAVY
        heading_style.font.name = 'Calibri'
        if i == 1:
            heading_style.font.size = Pt(20)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif i == 2:
            heading_style.font.size = Pt(16)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            heading_style.font.size = Pt(13)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    return doc


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with navy header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A5F"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_severity_text(paragraph, level):
    """Add colored severity label to a paragraph."""
    color = SEVERITY_COLORS.get(level, MED_GRAY)
    label = SEVERITY_LABELS.get(level, level)
    run = paragraph.add_run(label)
    run.font.color.rgb = color
    run.bold = True
    run.font.size = Pt(10)
    return run


def add_key_value(doc, key, value, bold_value=False):
    """Add a 'Key: Value' line."""
    p = doc.add_paragraph()
    run_k = p.add_run(f"{key}: ")
    run_k.bold = True
    run_k.font.size = Pt(11)
    run_k.font.color.rgb = DARK_GRAY
    run_v = p.add_run(str(value))
    run_v.font.size = Pt(11)
    run_v.font.color.rgb = DARK_GRAY
    if bold_value:
        run_v.bold = True
    return p


def add_body_text(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
    return p


def add_bullet(doc, text):
    """Add a bulleted list item."""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
    return p


# =============================================================================
# Document Sections
# =============================================================================

def add_cover_page(doc):
    """Add the cover page."""
    for _ in range(6):
        doc.add_paragraph()

    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Application Security\nAssessment Report")
    run.font.size = Pt(32)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cadence \u2014 HSEEP MSEL Management Platform")
    run.font.size = Pt(16)
    run.font.color.rgb = MED_GRAY
    run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    meta_items = [
        ("Report Date", datetime.now().strftime("%B %d, %Y")),
        ("Scan Date", SCAN_DATE),
        ("Document Version", "1.3"),
        ("Classification", "CONFIDENTIAL \u2014 For Authorized Recipients Only"),
        ("Assessment Type", "Internal Automated Security Assessment"),
        ("Prepared By", "Dynamis Inc. \u2014 Security Engineering"),
    ]
    for key, val in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_k = p.add_run(f"{key}:  ")
        run_k.bold = True
        run_k.font.size = Pt(11)
        run_k.font.color.rgb = DARK_GRAY
        run_v = p.add_run(val)
        run_v.font.size = Pt(11)
        run_v.font.color.rgb = DARK_GRAY

    doc.add_page_break()


def add_headers_footers(doc):
    """Set up header/footer."""
    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False
    h_para = header.paragraphs[0]
    h_para.text = ""
    run_left = h_para.add_run("Cadence \u2014 Security Assessment")
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = MED_GRAY
    run_left.font.name = 'Calibri'
    h_para.add_run("\t\t")
    run_right = h_para.add_run("CONFIDENTIAL")
    run_right.font.size = Pt(8)
    run_right.font.color.rgb = RED
    run_right.bold = True
    run_right.font.name = 'Calibri'

    footer = section.footer
    footer.is_linked_to_previous = False
    f_para = footer.paragraphs[0]
    f_para.text = ""
    run_left = f_para.add_run("Dynamis Inc.")
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = MED_GRAY
    run_left.font.name = 'Calibri'
    f_para.add_run("\t\t")
    run_right = f_para.add_run("Page ")
    run_right.font.size = Pt(8)
    run_right.font.color.rgb = MED_GRAY

    run_begin = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="begin"/></w:r>')
    run_instr = parse_xml(f'<w:r {nsdecls("w")}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    run_end = parse_xml(f'<w:r {nsdecls("w")}><w:fldChar w:fldCharType="end"/></w:r>')
    f_para._p.append(run_begin)
    f_para._p.append(run_instr)
    f_para._p.append(run_end)


def add_table_of_contents(doc):
    """Add the table of contents page."""
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph("Update this field in Word: right-click \u2192 Update Field, or press Ctrl+A then F9")
    p.italic = True
    p.runs[0].font.color.rgb = MED_GRAY
    p.runs[0].font.size = Pt(9)

    p_toc = doc.add_paragraph()
    run = p_toc.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    doc.add_page_break()


def add_executive_summary(doc, tools, deduped, all_findings):
    """Section 1: Executive Summary."""
    doc.add_heading("1. Executive Summary", level=1)

    total_instances = len(all_findings)
    unique_findings = len(deduped)
    sast_findings = [f for f in all_findings if "Semgrep" in f["tool"]]
    dast_findings = [f for f in all_findings if "ZAP" in f["tool"]]
    iac_findings = [f for f in all_findings if "Checkov" in f["tool"]]

    # Self-assessment disclaimer
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = ORANGE
    run.font.size = Pt(11)
    run = p.add_run(
        "This report is an internal automated security assessment conducted by the vendor (Dynamis Inc.), "
        "not an independent third-party audit. While the automated tooling provides objective, reproducible "
        "results, the risk analysis and contextual interpretation are performed by the development team. "
        "Organizations with stringent compliance requirements should consider supplementing this assessment "
        "with an independent penetration test or third-party security audit."
    )
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacing

    add_body_text(doc,
        "Dynamis Inc. maintains a continuous automated security assessment program for the Cadence platform. "
        "This report provides a transparent view into the current security posture of the application, "
        "infrastructure, and live deployment environment. It is produced as part of an ongoing commitment to "
        "proactive vulnerability identification and remediation\u2014not as a one-time audit, but as evidence of "
        "continuous security vigilance."
    )

    add_body_text(doc,
        "The assessment employs a defense-in-depth strategy using three independent and complementary security tools, "
        "each targeting a different layer of the application stack. Static Application Security Testing (SAST) examines "
        "the source code for vulnerabilities at the code level. Infrastructure as Code (IaC) scanning validates Azure "
        "resource configurations against industry benchmarks before deployment. Dynamic Application Security Testing (DAST) "
        "probes the live UAT environment to identify runtime vulnerabilities that static analysis cannot detect."
    )

    add_body_text(doc,
        "These automated scans run on a defined schedule integrated into the CI/CD pipeline. Every code change to the "
        "main branch triggers SAST analysis. Infrastructure templates are scanned on each relevant commit. DAST baseline "
        "scans execute weekly, with comprehensive API-level scans running monthly. Results are consolidated into a unified "
        "SARIF report weekly and monthly, with findings automatically tracked via GitHub Issues."
    )

    doc.add_heading("Key Results", level=3)

    # SAST results — tempered claims acknowledging tooling limitations
    p = doc.add_paragraph()
    run = p.add_run("Static Application Security Testing (SAST): ")
    run.bold = True
    run.font.color.rgb = DARK_GRAY
    run = p.add_run("No findings detected within configured rule packs. ")
    run.font.color.rgb = GREEN
    run.bold = True
    run = p.add_run(
        "Semgrep OSS analyzed the C# and TypeScript codebase against the OWASP Top 10 2025, CWE rules, "
        "and secret detection rule packs at ERROR and WARNING severity thresholds. "
        "Note: Semgrep OSS has limited depth for C# compared to commercial SAST tools. "
        "The TypeScript/React rule coverage is more comprehensive. These results should be "
        "interpreted as evidence that no issues were detected within the scope of the configured "
        "rules, not as a guarantee that zero vulnerabilities exist in the codebase."
    )
    run.font.color.rgb = DARK_GRAY

    p = doc.add_paragraph()
    run = p.add_run("Dynamic Application Security Testing (DAST): ")
    run.bold = True
    dast_unique = len(set(f['rule_id'] for f in dast_findings))
    run = p.add_run(
        f"{len(dast_findings)} instances across {dast_unique} unique rules. "
        "Findings relate to HTTP header configuration recommendations (HSTS, caching directives). "
        "No injection vulnerabilities, cross-site scripting, or authentication bypasses were detected "
        "during the unauthenticated baseline scan or the authenticated OpenAPI-driven API scan. "
        "Note: DAST coverage is limited to HTTP REST endpoints. SignalR WebSocket connections, "
        "file upload handlers, and complex multi-step workflows are not covered by the current "
        "automated DAST configuration and would require manual testing or specialized tooling."
    )

    p = doc.add_paragraph()
    run = p.add_run("Infrastructure as Code (IaC): ")
    run.bold = True
    run = p.add_run(
        f"{len(iac_findings)} instances across {len(set(f['rule_id'] for f in iac_findings))} unique rules. "
        "Checkov identified configuration hardening opportunities in the Azure Bicep templates. "
        "These findings represent best-practice recommendations from CIS benchmarks\u2014the majority relate to "
        "enabling additional Azure security features (Defender, auditing, zone redundancy) and network isolation "
        "settings appropriate for enterprise-grade deployments. Many are cost-tier dependent and are evaluated "
        "against the platform's current deployment model."
    )

    doc.add_heading("Summary", level=3)

    iac_unique = len(set(f['rule_id'] for f in iac_findings))
    medium_count = len([f for f in deduped if RISK_CLASSIFICATION.get(f['rule_id']) == 'Medium'])
    low_count = len([f for f in deduped if RISK_CLASSIFICATION.get(f['rule_id']) == 'Low'])
    info_count = len([f for f in deduped if RISK_CLASSIFICATION.get(f['rule_id'], 'Informational') == 'Informational'])

    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Total Scan Instances", str(total_instances)],
            ["Unique Findings (Deduplicated)", str(unique_findings)],
            ["SAST Findings (Source Code)", "0 (within configured rule packs)"],
            ["DAST Findings (Runtime)", f"{dast_unique} header recommendations"],
            ["IaC Findings (Infrastructure)", f"{iac_unique} configuration recommendations"],
            ["Assessed Risk: Medium", str(medium_count)],
            ["Assessed Risk: Low", str(low_count)],
            ["Assessed Risk: Informational", str(info_count)],
            ["Scan Tools", "3 (Semgrep OSS, OWASP ZAP, Checkov)"],
            ["Scan Pipeline", "Automated CI/CD (continuous)"],
        ],
        col_widths=[3.5, 3.0]
    )

    add_body_text(doc,
        "The assessment indicates an actively maintained security posture with a continuous automated scanning pipeline. "
        "No exploitable source code or runtime vulnerabilities were detected within the scope of the configured tools. "
        f"{medium_count} findings are classified as Medium risk, relating primarily to infrastructure hardening "
        "opportunities (network isolation, audit logging, managed identity). These have specific remediation plans "
        "documented in Section 7. Risk ratings are calibrated against the threat model in Section 2.6, which considers "
        "adversary profiles relevant to a multi-tenant emergency management platform handling sensitive HSEEP exercise data. "
        "The absence of Critical or High risk findings is encouraging, though this assessment "
        "should be supplemented with manual penetration testing to validate coverage beyond automated tooling."
    )

    doc.add_page_break()


def add_scope_methodology(doc):
    """Section 2: Scope & Methodology."""
    doc.add_heading("2. Scope & Methodology", level=1)

    doc.add_heading("2.1 Assessment Scope", level=2)
    add_body_text(doc,
        "The following table defines the components and layers included in this security assessment."
    )

    add_styled_table(doc,
        ["Component", "Technology", "Scan Coverage"],
        [
            ["Backend REST API", ".NET 10 / ASP.NET Core", "SAST (Semgrep), DAST (ZAP API scan)"],
            ["Frontend SPA", "React 19 / TypeScript 5", "SAST (Semgrep)"],
            ["Infrastructure", "Azure Bicep (IaC)", "IaC (Checkov)"],
            ["Live UAT Environment", "Azure App Service", "DAST (ZAP baseline + API scan)"],
            ["Authentication", "JWT Bearer / Microsoft Entra ID", "SAST patterns, DAST authentication flows"],
            ["Real-Time Layer", "Azure SignalR Service", "Architecture review (not directly scanned)"],
            ["Database", "Azure SQL", "IaC configuration checks"],
        ],
        col_widths=[2.0, 2.0, 2.8]
    )

    doc.add_heading("2.2 Security Testing Tools", level=2)

    add_styled_table(doc,
        ["Tool", "Version", "Type", "Standards"],
        [
            ["Semgrep OSS", "1.153.1", "SAST", "OWASP Top 10 2025, CWE, SARIF 2.1.0"],
            ["OWASP ZAP", "2.17.0", "DAST", "OWASP Testing Guide, SARIF 2.1.0"],
            ["Checkov", "3.2.507", "IaC", "CIS Benchmarks, Azure Best Practices, SARIF 2.1.0"],
        ],
        col_widths=[1.8, 1.2, 1.0, 2.8]
    )

    doc.add_heading("2.3 Scan Configuration", level=2)

    doc.add_heading("Semgrep (SAST)", level=3)
    add_body_text(doc,
        "Semgrep runs against the full source tree with the following rule packs enabled: auto, csharp, "
        "typescript, react, secrets, and owasp-top-ten. Only findings at ERROR and WARNING severity are included "
        "in the consolidated report. The scan covers C# backend code, TypeScript/React frontend code, and "
        "configuration files for hardcoded secrets."
    )

    doc.add_heading("OWASP ZAP (DAST)", level=3)
    add_body_text(doc,
        "ZAP runs two scan profiles against the UAT environment:"
    )
    add_bullet(doc, "Baseline Scan (weekly): Unauthenticated passive scan that spiders the application and checks HTTP headers, cookies, and common misconfigurations.")
    add_bullet(doc, "API Scan (monthly): Authenticated active scan driven by the OpenAPI specification. Tests every API endpoint for injection, XSS, authentication bypass, and other OWASP attack vectors.")

    add_body_text(doc,
        "ZAP scan behavior is customized via a rules.tsv configuration file that excludes known false positives "
        "and accepted risks:"
    )
    add_styled_table(doc,
        ["Rule ID", "Action", "Rationale"],
        [
            ["10038", "IGNORE", "CSP Header Not Set \u2014 will address with CSP implementation"],
            ["10063", "IGNORE", "Permissions Policy Header Not Set \u2014 low risk for internal API"],
            ["90033", "IGNORE", "Loosely Scoped Cookie \u2014 Azure auth cookies are scoped by platform"],
            ["10202", "IGNORE", "Anti-CSRF Tokens \u2014 API uses JWT bearer auth, not cookie-based sessions"],
            ["10098", "WARN", "Cross-Domain Misconfiguration \u2014 review if flagged"],
            ["10020", "WARN", "X-Frame-Options Header \u2014 should be set but not blocking"],
            ["10021", "WARN", "X-Content-Type-Options Header Missing \u2014 should be set"],
            ["10036", "WARN", "Server Leaks Version Information \u2014 cosmetic but worth fixing"],
            ["40012", "WARN", "Reflected XSS \u2014 review findings from active scan"],
            ["40014", "WARN", "Persistent XSS \u2014 review findings from active scan"],
            ["40018", "WARN", "SQL Injection \u2014 review findings from active scan"],
            ["40003", "WARN", "CRLF Injection \u2014 review findings from active scan"],
        ],
        col_widths=[1.0, 1.0, 4.8]
    )

    doc.add_heading("Checkov (IaC)", level=3)
    add_body_text(doc,
        "Checkov scans all Azure Bicep templates in the infrastructure/ directory against the CIS Azure "
        "benchmark and Checkov's built-in Azure policy framework. The scan evaluates resource configurations "
        "for security best practices including encryption, network isolation, identity management, and logging."
    )

    doc.add_heading("2.4 Limitations & Coverage Gaps", level=2)
    add_body_text(doc,
        "This assessment has the following known limitations that should be considered when evaluating the results:"
    )
    add_bullet(doc, "No manual penetration testing: This is an automated-only assessment. Manual testing by a skilled tester can find business logic flaws, authorization bypasses, and chained vulnerabilities that automated tools miss.")
    add_bullet(doc, "No red team or social engineering testing.")
    add_bullet(doc, "SAST C# coverage is limited: Semgrep OSS has fewer rules for C# compared to TypeScript or Python. Commercial SAST tools (e.g., Semgrep Pro, Snyk Code, SonarQube) offer deeper C#/.NET analysis including taint tracking and data flow analysis.")
    add_bullet(doc, "DAST does not cover SignalR WebSocket connections, file upload handlers, or complex multi-step authentication flows. These would require specialized testing tools or manual assessment.")
    add_bullet(doc, "DAST scan uses a single authenticated user context. It does not test horizontal privilege escalation (accessing another user's data) or vertical privilege escalation (performing admin actions as a regular user).")
    add_bullet(doc, "Dependency vulnerability scanning (SCA) is handled separately by GitHub Dependabot and is not included in this report.")
    add_bullet(doc, "The DAST scan targets the UAT environment, which mirrors production configuration but may differ in data volume, network topology, and load characteristics.")
    add_bullet(doc, "IaC scanning evaluates Bicep templates only. Runtime Azure configuration drift from the templates is not detected.")

    doc.add_heading("2.5 Architecture Review Scope", level=2)
    add_body_text(doc,
        "The security architecture described in Section 3 was reviewed by the development team against the "
        "OWASP Application Security Verification Standard (ASVS) Level 1 requirements. The review examined "
        "authentication flows, authorization enforcement, multi-tenancy isolation, and data protection controls "
        "as implemented in the source code and infrastructure templates."
    )
    add_body_text(doc,
        "This architecture review was conducted internally by the Dynamis engineering team, not by an "
        "independent reviewer. The review methodology consisted of code inspection (examining authentication "
        "middleware, authorization attributes, EF Core query filters, and JWT claim handling) and infrastructure "
        "template analysis. No formal architecture threat modeling workshop was conducted. The threat model in "
        "Section 2.6 was developed from the architecture review and should be validated by an independent "
        "security professional."
    )

    doc.add_heading("2.6 Threat Model", level=2)
    add_body_text(doc,
        "The risk ratings in this report are calibrated against the following threat model. Without a threat "
        "model, risk classifications would be the assessor's intuition rather than a structured analysis."
    )

    doc.add_heading("Adversary Profiles", level=3)
    add_styled_table(doc,
        ["Adversary", "Motivation", "Capability", "Likelihood"],
        [
            ["Unauthorized External Actor", "Access exercise data (vulnerability scenarios, critical infrastructure details) for intelligence or disruption", "Automated scanning, credential stuffing, known CVE exploitation", "Medium"],
            ["Authenticated Insider (Cross-Tenant)", "Access another organization's exercise data through privilege escalation or tenant isolation bypass", "Valid credentials, API knowledge, patience to probe authorization boundaries", "Medium"],
            ["Malicious Exercise Participant", "Modify inject data, disrupt exercise conduct, or exfiltrate MSEL content", "Authenticated session with Observer/Evaluator role, browser developer tools", "Low-Medium"],
            ["Compromised CI/CD Pipeline", "Supply chain attack via dependency poisoning or workflow manipulation", "Access to package registries or GitHub Actions marketplace", "Low"],
            ["Opportunistic Attacker", "Ransomware, cryptomining, or generic exploitation of exposed services", "Automated scanning, commodity attack tools", "Medium"],
        ],
        col_widths=[1.5, 2.2, 2.0, 1.1]
    )

    doc.add_heading("Impact Analysis", level=3)
    add_body_text(doc,
        "Compromise of Cadence could result in:"
    )
    add_bullet(doc, "Confidentiality: Exposure of HSEEP exercise scenarios, vulnerability assessments, critical infrastructure details, and organizational response capabilities. This data could inform real-world attack planning.")
    add_bullet(doc, "Integrity: Modification of MSEL injects during live exercise conduct, potentially causing incorrect emergency response actions or undermining exercise validity.")
    add_bullet(doc, "Availability: Disruption of exercise conduct operations during time-critical training events, which may have regulatory or compliance deadlines.")
    add_bullet(doc, "Multi-tenant breach: Cross-organization data exposure affecting multiple emergency management agencies using the shared platform.")

    add_body_text(doc,
        "These impact scenarios inform the risk classifications in Section 6. Findings that could enable "
        "cross-tenant data access or exercise disruption are rated Medium or higher. Findings limited to "
        "defense-in-depth hardening are rated Low or Informational."
    )

    doc.add_heading("2.7 Data Classification", level=2)
    add_body_text(doc,
        "Cadence processes the following categories of data. Risk ratings in this report are calibrated "
        "to the sensitivity of the most critical data category handled by the system."
    )
    add_styled_table(doc,
        ["Data Category", "Classification", "Examples", "Protection Requirements"],
        [
            ["Exercise Scenarios (MSEL)", "Sensitive / Official Use Only", "Inject descriptions, scenario narratives, expected responses, vulnerability details", "Encryption at rest and in transit, organization-scoped access control, audit logging"],
            ["Organizational Structure", "Sensitive", "Agency names, team structures, capability assessments, resource inventories", "Tenant isolation, role-based access"],
            ["User Credentials", "Confidential", "Email addresses, password hashes, JWT signing keys", "Hashing (bcrypt), encrypted storage, Key Vault (planned)"],
            ["Exercise Observations", "Sensitive / Official Use Only", "Performance assessments, improvement plans, identified weaknesses", "Organization-scoped access, soft delete for retention"],
            ["System Configuration", "Internal", "Connection strings, API keys, deployment parameters", "App Service settings (current), Key Vault (planned)"],
            ["Audit & Telemetry", "Internal", "Access logs, scan results, error traces", "Azure platform logging, 90-day retention"],
        ],
        col_widths=[1.5, 1.3, 2.2, 1.8]
    )

    add_body_text(doc,
        "Note: HSEEP exercises may involve scenarios based on real-world threat assessments and critical "
        "infrastructure vulnerabilities. While exercise data is typically marked For Official Use Only (FOUO) "
        "or equivalent, Cadence does not currently enforce formal data classification labels within the "
        "application. Organizations handling classified or law-enforcement-sensitive exercise content should "
        "evaluate whether additional controls beyond those described in this report are required."
    )

    doc.add_page_break()


def add_security_architecture(doc):
    """Section 3: Security Architecture Overview."""
    doc.add_heading("3. Security Architecture Overview", level=1)

    add_body_text(doc,
        "Cadence implements a layered security architecture designed for multi-tenant emergency management "
        "operations. The following sections describe the key security controls in place."
    )

    doc.add_heading("3.1 Authentication", level=2)
    add_body_text(doc,
        "Cadence uses JWT (JSON Web Token) bearer authentication. Tokens are issued upon successful "
        "authentication and include claims for user identity, system role, organization context, and "
        "organization-level role. Tokens are validated on every API request. The authentication system "
        "is designed for integration with Microsoft Entra ID (formerly Azure AD) for enterprise single sign-on (SSO) scenarios."
    )

    doc.add_heading("3.2 Authorization \u2014 Three-Tier Role Hierarchy", level=2)
    add_body_text(doc,
        "Cadence enforces authorization through a three-tier role model that provides granular access control:"
    )
    add_styled_table(doc,
        ["Tier", "Role Type", "Scope", "Examples"],
        [
            ["System", "SystemRole", "Platform-wide", "Admin, Manager, User"],
            ["Organization", "OrgRole", "Per-organization", "OrgAdmin, OrgManager, OrgUser"],
            ["Exercise", "ExerciseRole", "Per-exercise", "ExerciseDirector, Controller, Evaluator, Observer"],
        ],
        col_widths=[1.2, 1.5, 1.5, 2.6]
    )
    add_body_text(doc,
        "Exercise roles follow HSEEP (Homeland Security Exercise and Evaluation Program) standards. "
        "Each role carries specific permissions that control what actions a user can perform within an exercise context."
    )

    doc.add_heading("3.3 Multi-Tenancy Data Isolation", level=2)
    add_body_text(doc,
        "Organization is the primary security and data isolation boundary. All domain entities (exercises, "
        "injects, observations) are scoped to an organization via an OrganizationId foreign key. Data isolation "
        "is enforced at multiple levels:"
    )
    add_bullet(doc, "Read-side: All database queries are filtered by the current user's organization context, extracted from the JWT token.")
    add_bullet(doc, "Write-side: An OrganizationValidationInterceptor ensures that entities are always saved with the correct organization ID, preventing cross-tenant data leakage.")
    add_bullet(doc, "API-level: Organization context is derived from JWT claims and injected into services via ICurrentOrganizationContext, eliminating the possibility of client-side organization spoofing.")

    doc.add_heading("3.4 Token & Session Management", level=2)
    add_body_text(doc,
        "Cadence uses stateless JWT bearer tokens for API authentication. Key characteristics:"
    )
    add_bullet(doc, "Token issuance: Tokens are issued by the Cadence API upon successful authentication with email/password credentials.")
    add_bullet(doc, "Token contents: Claims include user identity (sub), email, system role, organization context (org_id), and organization role (org_role).")
    add_bullet(doc, "Token validation: Every API request validates the JWT signature, expiration, and issuer claims.")
    add_bullet(doc, "Token storage: Tokens are stored client-side (browser memory/localStorage). No server-side session state is maintained.")
    add_bullet(doc, "Token revocation: Token expiry is the primary revocation mechanism. Refresh token rotation and server-side session invalidation are not currently implemented.")
    add_bullet(doc, "MFA: Not enforced at the application level. Can be configured at the identity provider level when using Microsoft Entra ID integration.")

    doc.add_heading("3.5 Data Protection", level=2)
    add_bullet(doc, "Encryption at rest: Azure SQL Transparent Data Encryption (TDE) is enabled by default for all databases.")
    add_bullet(doc, "Encryption in transit: All communication uses HTTPS/TLS. Azure App Service enforces TLS 1.2 minimum.")

    add_bullet(doc, "Secrets management: Sensitive configuration values (JWT keys, connection strings) are stored as App Service application settings, not in source code. Semgrep secret detection rules verify no hardcoded secrets exist in the codebase. Azure Key Vault is not currently used; secrets are not versioned or rotation-audited. See Section 7.3 for remediation plan.")

    doc.add_heading("3.6 API Security", level=2)
    add_bullet(doc, "CORS: Configured to allow only the specific frontend domain, preventing unauthorized cross-origin requests.")
    add_bullet(doc, "Input validation: Request models are validated using data annotations and FluentValidation. Entity Framework Core parameterizes all queries, preventing SQL injection by design.")
    add_bullet(doc, "Rate limiting: Azure App Service provides built-in request throttling capabilities.")
    add_bullet(doc, "CSRF protection: Not applicable\u2014the API uses JWT bearer authentication (not cookies), so CSRF attacks are not a viable vector.")

    doc.add_heading("3.7 Deployment Security", level=2)
    add_body_text(doc,
        "The application runs on Azure App Service (B1 tier), a fully managed platform that handles OS patching, "
        "runtime updates, and infrastructure security. The frontend is deployed as an Azure Static Web App with "
        "built-in CDN and DDoS protection. Infrastructure is defined as code (Bicep templates) and deployed through "
        "automated CI/CD pipelines, ensuring reproducible and auditable deployments."
    )

    doc.add_page_break()


def add_findings_summary(doc, deduped, all_findings):
    """Section 4: Findings Summary."""
    doc.add_heading("4. Findings Summary", level=1)

    doc.add_heading("4.1 Overview Metrics", level=2)

    add_body_text(doc,
        "Findings are classified at two levels: the raw tool severity (what the scanning tool reported) "
        "and the assessed risk (contextual classification considering the application architecture, existing "
        "mitigations, and real-world exploitability). The assessed risk is the primary classification used "
        "throughout this report."
    )

    # Table 1: By assessed risk (PRIMARY)
    p = doc.add_paragraph()
    run = p.add_run("Findings by Assessed Risk")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    risk_counts_unique = Counter(RISK_CLASSIFICATION.get(f["rule_id"], "Informational") for f in deduped)
    risk_counts_total = Counter(RISK_CLASSIFICATION.get(f["rule_id"], "Informational") for f in all_findings)

    add_styled_table(doc,
        ["Assessed Risk", "Unique Findings", "Total Instances"],
        [
            ["Critical", str(risk_counts_unique.get("Critical", 0)), str(risk_counts_total.get("Critical", 0))],
            ["High", str(risk_counts_unique.get("High", 0)), str(risk_counts_total.get("High", 0))],
            ["Medium", str(risk_counts_unique.get("Medium", 0)), str(risk_counts_total.get("Medium", 0))],
            ["Low", str(risk_counts_unique.get("Low", 0)), str(risk_counts_total.get("Low", 0))],
            ["Informational", str(risk_counts_unique.get("Informational", 0)), str(risk_counts_total.get("Informational", 0))],
            ["Total", str(len(deduped)), str(len(all_findings))],
        ],
        col_widths=[2.0, 2.0, 2.0]
    )

    # Table 2: By tool severity (for transparency)
    p = doc.add_paragraph()
    run = p.add_run("Findings by Tool-Reported Severity (Raw)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY

    add_body_text(doc,
        "The table below shows the raw severity levels reported by each tool. Note that Checkov reports "
        "all IaC misconfigurations as SARIF 'error' regardless of actual risk. These are re-assessed in "
        "the Assessed Risk classification above based on exploitability, blast radius, and existing controls."
    )

    level_counts = Counter(f["level"] for f in all_findings)
    unique_by_level = Counter(f["level"] for f in deduped)

    add_styled_table(doc,
        ["Tool Severity", "Total Instances", "Unique Findings", "Note"],
        [
            ["Error", str(level_counts.get("error", 0)), str(unique_by_level.get("error", 0)),
             "Checkov reports all IaC findings as error"],
            ["Warning", str(level_counts.get("warning", 0)), str(unique_by_level.get("warning", 0)),
             "ZAP Low/Medium risk findings"],
            ["Note", str(level_counts.get("note", 0)), str(unique_by_level.get("note", 0)),
             "ZAP informational findings"],
            ["Total", str(len(all_findings)), str(len(deduped)), ""],
        ],
        col_widths=[1.3, 1.3, 1.3, 2.9]
    )

    doc.add_heading("4.2 Findings by Tool", level=2)

    tool_groups = defaultdict(list)
    for f in all_findings:
        tool_groups[f["tool"]].append(f)

    rows = []
    for tool_name in ["Semgrep OSS", "OWASP ZAP", "Checkov"]:
        findings = tool_groups.get(tool_name, [])
        unique = len(set(f["rule_id"] for f in findings))
        rows.append([tool_name, str(len(findings)), str(unique)])
    rows.append(["Total", str(len(all_findings)), str(len(deduped))])

    add_styled_table(doc,
        ["Tool", "Instances", "Unique Rules"],
        rows,
        col_widths=[2.0, 2.0, 2.0]
    )

    doc.add_heading("4.3 OWASP Top 10 2025 Mapping", level=2)
    add_body_text(doc,
        "The following table maps each finding to its corresponding OWASP Top 10 2025 category. "
        "Findings are mapped via CWE identifiers where available."
    )

    owasp_groups = defaultdict(list)
    for f in deduped:
        rule_id = f["rule_id"]
        if rule_id in CHECKOV_OWASP:
            owasp_id, owasp_name, _ = CHECKOV_OWASP[rule_id]
            owasp_groups[(owasp_id, owasp_name)].append(rule_id)
        elif f["tags"]:
            for tag in f["tags"]:
                if tag in CWE_TO_OWASP:
                    owasp_id, owasp_name = CWE_TO_OWASP[tag]
                    owasp_groups[(owasp_id, owasp_name)].append(rule_id)
                    break

    owasp_rows = []
    for (owasp_id, owasp_name), rule_ids in sorted(owasp_groups.items()):
        unique_rules = sorted(set(rule_ids))
        owasp_rows.append([owasp_id, owasp_name, ", ".join(unique_rules), str(len(unique_rules))])

    # Add categories with 0 findings (OWASP Top 10 2025)
    all_owasp = [
        ("A01:2025", "Broken Access Control"),
        ("A02:2025", "Security Misconfiguration"),
        ("A03:2025", "Software Supply Chain Failures"),
        ("A04:2025", "Cryptographic Failures"),
        ("A05:2025", "Injection"),
        ("A06:2025", "Insecure Design"),
        ("A07:2025", "Authentication Failures"),
        ("A08:2025", "Software or Data Integrity Failures"),
        ("A09:2025", "Logging & Alerting Failures"),
        ("A10:2025", "Mishandling of Exceptional Conditions"),
    ]
    seen_owasp = set(k[0] for k in owasp_groups.keys())
    for owasp_id, owasp_name in all_owasp:
        if owasp_id not in seen_owasp:
            owasp_rows.append([owasp_id, owasp_name, "None", "0"])

    owasp_rows.sort(key=lambda x: x[0])

    add_styled_table(doc,
        ["OWASP ID", "Category", "Related Findings", "Count"],
        owasp_rows,
        col_widths=[1.2, 2.5, 2.0, 0.8]
    )

    doc.add_page_break()


def add_detailed_findings(doc, deduped):
    """Section 5: Detailed Findings."""
    doc.add_heading("5. Detailed Findings", level=1)

    add_body_text(doc,
        "This section provides a detailed analysis of each unique finding, grouped by assessed risk level. "
        "Findings are deduplicated by rule ID\u2014instance counts reflect the number of occurrences across "
        "all scanned targets. The assessed risk reflects real-world exploitability in the context of the "
        "Cadence platform, not the raw severity reported by the scanning tool."
    )

    # Re-sort deduped by assessed risk instead of SARIF level
    deduped_by_risk = sorted(deduped, key=lambda x: (
        RISK_ORDER.get(RISK_CLASSIFICATION.get(x["rule_id"], "Informational"), 4),
        x["rule_id"]
    ))

    current_risk = None
    risk_section_idx = 0

    for f in deduped_by_risk:
        risk = RISK_CLASSIFICATION.get(f["rule_id"], "Informational")
        if risk != current_risk:
            current_risk = risk
            risk_section_idx += 1
            doc.add_heading(f"5.{risk_section_idx} {risk.upper()} Risk Findings", level=2)

        rule_id = f["rule_id"]
        short = f["short_desc"]

        # Finding header
        doc.add_heading(f"{rule_id}: {short}", level=3)

        # Metadata table
        level = f["level"]
        tool = f["tool"]
        tags = f["tags"]
        cwe = tags[0] if tags else ""

        # Get OWASP mapping
        owasp_str = "N/A"
        if rule_id in CHECKOV_OWASP:
            owasp_id, owasp_name, cwe_from_map = CHECKOV_OWASP[rule_id]
            owasp_str = f"{owasp_id} ({owasp_name})"
            if not cwe:
                cwe = cwe_from_map
        elif cwe and cwe in CWE_TO_OWASP:
            owasp_id, owasp_name = CWE_TO_OWASP[cwe]
            owasp_str = f"{owasp_id} ({owasp_name})"

        risk_class = RISK_CLASSIFICATION.get(rule_id, "Informational")
        meta_rows = [
            ["Assessed Risk", risk_class],
            ["Tool Severity", TOOL_SEVERITY_LABELS.get(level, level)],
            ["Tool", tool],
            ["Instances", str(f["instance_count"])],
            ["CWE", cwe if cwe else "N/A"],
            ["OWASP Category", owasp_str],
        ]
        if f.get("confidence"):
            conf = f["confidence"]
            conf_label = "Confirmed" if conf >= 99 else "High" if conf >= 90 else "Medium" if conf >= 60 else "Low"
            meta_rows.append(["Confidence", f"{conf_label} ({conf})"])

        add_styled_table(doc, ["Property", "Value"], meta_rows, col_widths=[1.8, 5.0])

        # Description
        p = doc.add_paragraph()
        run = p.add_run("Description: ")
        run.bold = True
        run.font.color.rgb = DARK_GRAY

        desc = ""
        if rule_id in CHECKOV_DESCRIPTIONS:
            desc = CHECKOV_DESCRIPTIONS[rule_id]
        elif f["full_desc"]:
            desc = f["full_desc"][:500]
        else:
            desc = f["message"].split("|")[0].strip()
        run = p.add_run(desc)
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(11)

        # Affected locations
        locs = f["unique_locations"]
        if locs:
            p = doc.add_paragraph()
            run = p.add_run("Affected Locations:")
            run.bold = True
            run.font.color.rgb = DARK_GRAY

            shown = locs[:5]
            for loc in shown:
                loc_text = ""
                if loc["method"]:
                    loc_text = f"{loc['method']} {loc['uri']}"
                elif loc["line"]:
                    loc_text = f"{loc['uri']}:{loc['line']}"
                else:
                    loc_text = loc["uri"]
                add_bullet(doc, loc_text)

            remaining = len(locs) - 5
            if remaining > 0:
                add_bullet(doc, f"...and {remaining} more location(s)")

        # Risk assessment
        p = doc.add_paragraph()
        run = p.add_run("Risk Assessment: ")
        run.bold = True
        run.font.color.rgb = DARK_GRAY

        risk_text = _get_risk_assessment(rule_id, f)
        run = p.add_run(risk_text)
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(11)

        # Remediation
        p = doc.add_paragraph()
        run = p.add_run("Remediation: ")
        run.bold = True
        run.font.color.rgb = DARK_GRAY

        remediation = _get_remediation(rule_id, f)
        run = p.add_run(remediation)
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(11)

    doc.add_page_break()


def _get_risk_assessment(rule_id, finding):
    """Return contextual risk assessment for a finding."""
    assessments = {
        "zap-10035": (
            "Mitigated. The Strict-Transport-Security (HSTS) header is now configured in ASP.NET Core middleware "
            "with max-age=31536000 (1 year) and includeSubDomains. This prevents SSL stripping attacks and protocol "
            "downgrade attacks on first connection."
        ),
        "zap-10049": (
            "Informational. Content caching behavior is flagged because responses include an Authorization header requirement, "
            "which prevents caching by shared proxies. This is the expected and correct behavior for an authenticated API\u2014"
            "API responses containing user-specific data should not be cached by intermediaries. No action required."
        ),
        "zap-90005": (
            "Informational. The Sec-Fetch-* headers (Dest, Mode, Site, User) are request-side headers sent by browsers "
            "to indicate the intent of a fetch request. Their absence in ZAP's scan requests is expected because ZAP is "
            "not a browser and does not send these headers. This finding reflects the scan tool's request behavior, not a "
            "server vulnerability. The server does not need to require these headers. No action required."
        ),
        "CKV_AZURE_222": (
            "Medium. The App Service currently allows public network access. While all API endpoints require JWT authentication, "
            "restricting network access to known IP ranges or a virtual network would add an additional layer of defense. "
            "This is a recommended hardening step for production environments handling sensitive exercise data."
        ),
        "CKV_AZURE_113": (
            "Medium. The SQL Server currently allows public network access. While authentication is required and firewall rules "
            "limit access, implementing private endpoints would eliminate exposure of the database endpoint to the public internet. "
            "This is an important hardening step for multi-tenant data protection."
        ),
        "CKV_AZURE_35": (
            "Mitigated. The Storage Account networkAcls are now configured with defaultAction: 'Deny' and "
            "bypass: 'AzureServices' in the Bicep template. Only Azure-trusted services can access the storage account."
        ),
        "CKV_AZURE_23": (
            "Mitigated. SQL Server auditing is now enabled in the Bicep template with audit events sent to "
            "Log Analytics workspace. Covers authentication successes/failures and batch operations."
        ),
        "CKV_AZURE_24": (
            "Mitigated. SQL Server audit retention is now configured to 90 days in the Bicep template, "
            "meeting compliance requirements for post-incident investigation."
        ),
        "CKV_AZURE_25": (
            "Mitigated. SQL Server Advanced Threat Protection is now enabled with all alert types active "
            "and email notifications to database administrators configured in the Bicep template."
        ),
        "CKV2_AZURE_27": (
            "Medium. Microsoft Entra ID authentication is not enabled as the primary authentication mechanism for the SQL Server. "
            "While SQL authentication is configured, Entra ID authentication provides centralized identity management, "
            "conditional access, and eliminates the need for password-based SQL credentials."
        ),
        "CKV_AZURE_71": (
            "Mitigated. The App Service now has a SystemAssigned managed identity configured in the Bicep template. "
            "The identity is granted Key Vault Secrets User role for credential-free secret access."
        ),
        "CKV_AZURE_84": (
            "Mitigated. Defender for Storage Standard tier (DefenderForStorageV2) is enabled in defender.bicep. "
            "This provides threat detection for blob storage including anomalous access patterns, potential data "
            "exfiltration, and suspicious upload activity."
        ),
        "CKV_AZURE_87": (
            "Mitigated. Defender for Key Vault Standard tier is enabled in defender.bicep. This monitors Key Vault "
            "access for unusual patterns including unauthorized access attempts and suspicious secret retrieval activity."
        ),
        "CKV_AZURE_19": (
            "Mitigated. Defender for Cloud Standard pricing tier is enabled for SQL resources in defender.bicep. "
            "This provides advanced threat detection for Azure SQL databases including SQL injection detection, "
            "brute force attempts, and anomalous database access patterns."
        ),
    }
    if rule_id in assessments:
        return assessments[rule_id]

    # Default assessments by tool
    if "Checkov" in finding["tool"]:
        return (
            f"{'Low' if RISK_CLASSIFICATION.get(rule_id) == 'Low' else 'Informational'}. "
            "This is an infrastructure configuration recommendation from CIS benchmarks. "
            "The finding represents a best-practice hardening opportunity rather than an actively exploitable vulnerability. "
            "The risk should be evaluated against the platform's current deployment tier and cost constraints."
        )
    return "Informational. This finding represents a best-practice recommendation with minimal direct security impact."


def _get_remediation(rule_id, finding):
    """Return specific remediation steps for a finding."""
    remediations = {
        "zap-10035": (
            "Remediated. HSTS middleware (app.UseHsts()) added to Program.cs with max-age=31536000 and "
            "includeSubDomains. Applied in non-Development environments only."
        ),
        "zap-10049": (
            "No remediation required. The non-storable content behavior is correct for an authenticated API. "
            "Responses containing user-specific data should not be cached by shared proxies."
        ),
        "zap-90005": (
            "No remediation required. Sec-Fetch-* headers are request-side browser headers, not server configuration. "
            "The ZAP scanner does not send these headers because it is not a browser."
        ),
        "CKV_AZURE_222": (
            "Configure the App Service to restrict public network access. Add access restrictions in the Bicep template "
            "or configure an Azure Virtual Network integration with a private endpoint. "
            "Reference: https://learn.microsoft.com/en-us/azure/app-service/networking-features"
        ),
        "CKV_AZURE_113": (
            "Disable public network access on the SQL Server and configure a private endpoint. Update the Bicep template "
            "with publicNetworkAccess: 'Disabled' and add a privateEndpoints resource. "
            "Reference: https://learn.microsoft.com/en-us/azure/azure-sql/database/private-endpoint-overview"
        ),
        "CKV_AZURE_35": (
            "Remediated. Storage Account networkAcls configured with defaultAction: 'Deny' and bypass: 'AzureServices' "
            "in storage.bicep. Only Azure-trusted services (App Service, Functions) can access the storage account."
        ),
        "CKV_AZURE_23": (
            "Remediated. SQL Server auditing enabled in database.bicep with audit events sent to Log Analytics. "
            "Covers authentication successes/failures and batch operations with 90-day retention."
        ),
        "CKV_AZURE_24": (
            "Remediated. SQL Server audit retention set to 90 days in the auditingSettings resource."
        ),
        "CKV_AZURE_25": (
            "Remediated. SQL Server Advanced Threat Protection enabled with all alert types active and "
            "email notifications to database administrators in the securityAlertPolicies resource."
        ),
        "CKV2_AZURE_27": (
            "Enable Microsoft Entra ID authentication for the SQL Server. The Bicep template already accepts Entra ID admin "
            "parameters (sqlEntraAdminLogin, sqlEntraAdminObjectId)\u2014ensure these are configured in the parameter files. "
            "Reference: https://learn.microsoft.com/en-us/azure/azure-sql/database/authentication-aad-configure"
        ),
        "CKV_AZURE_71": (
            "Remediated. SystemAssigned managed identity enabled on the App Service in webapp.bicep. "
            "The identity is granted Key Vault Secrets User role via RBAC role assignment in keyvault.bicep."
        ),
        "CKV_AZURE_84": (
            "Remediated. Defender for Storage Standard tier (DefenderForStorageV2) enabled in defender.bicep "
            "via Microsoft.Security/defenderForStorageSettings resource. Provides real-time threat detection "
            "for blob storage operations."
        ),
        "CKV_AZURE_87": (
            "Remediated. Defender for Key Vault Standard tier enabled in defender.bicep via "
            "Microsoft.Security/pricings resource. Monitors vault access patterns and alerts on suspicious activity."
        ),
        "CKV_AZURE_19": (
            "Remediated. Defender for Cloud Standard pricing tier enabled for SqlServers in defender.bicep via "
            "Microsoft.Security/pricings resource. Provides SQL-specific threat detection including SQL injection "
            "and brute force alerts."
        ),
    }
    if rule_id in remediations:
        return remediations[rule_id]

    if finding.get("help_text"):
        return finding["help_text"][:400]

    if "Checkov" in finding["tool"]:
        return (
            "Review the CIS benchmark recommendation and update the corresponding Bicep template. "
            "See Checkov documentation for specific remediation guidance: https://www.checkov.io/"
        )
    return "Review the finding details and apply the recommended configuration change."


def add_risk_matrix(doc, deduped):
    """Section 6: Risk Assessment Matrix."""
    doc.add_heading("6. Risk Assessment Matrix", level=1)

    add_body_text(doc,
        "The following matrix classifies each finding based on its real-world risk in the context of the Cadence platform. "
        "Risk classification considers exploitability, blast radius, multi-tenancy implications, and existing mitigating controls. "
        "Many findings flagged as 'error' by IaC scanning tools represent best-practice configuration recommendations "
        "rather than actively exploitable vulnerabilities."
    )

    risk_levels = ["Critical", "High", "Medium", "Low", "Informational"]
    risk_descriptions = {
        "Critical": "Immediate action required. Actively exploitable vulnerability with high impact.",
        "High": "Address this sprint. Significant risk that should be remediated promptly.",
        "Medium": "Address this quarter. Important hardening that improves defense-in-depth.",
        "Low": "Monitor and plan. Best-practice improvement with minimal direct risk.",
        "Informational": "Accept or monitor. Informational findings with no direct security impact.",
    }

    rows = []
    for risk in risk_levels:
        matching = [f for f in deduped if RISK_CLASSIFICATION.get(f["rule_id"]) == risk]
        if matching:
            rule_ids = ", ".join(sorted(set(f["rule_id"] for f in matching)))
            count = str(len(matching))
        else:
            rule_ids = "None"
            count = "0"
        rows.append([risk, risk_descriptions[risk], count, rule_ids])

    add_styled_table(doc,
        ["Risk Level", "Definition", "Count", "Findings"],
        rows,
        col_widths=[1.2, 2.5, 0.6, 2.5]
    )

    critical_count = len([f for f in deduped if RISK_CLASSIFICATION.get(f["rule_id"]) == "Critical"])
    high_count = len([f for f in deduped if RISK_CLASSIFICATION.get(f["rule_id"]) == "High"])
    medium_count = len([f for f in deduped if RISK_CLASSIFICATION.get(f["rule_id"]) == "Medium"])

    if critical_count == 0 and high_count == 0:
        add_body_text(doc,
            f"No Critical or High risk findings were identified by the automated scanning tools. "
            f"{medium_count} Medium risk findings relate to infrastructure hardening and should be addressed "
            "per the remediation roadmap in Section 7. Note: the absence of Critical/High findings from automated "
            "tools does not guarantee the absence of all high-risk vulnerabilities\u2014manual penetration testing is "
            "recommended to validate this assessment."
        )
    else:
        add_body_text(doc,
            f"The assessment identified {critical_count} Critical and {high_count} High risk findings that "
            "require immediate attention. See the remediation roadmap in Section 7 for prioritized action items."
        )

    doc.add_page_break()


def add_remediation_roadmap(doc, deduped):
    """Section 7: Remediation Status & Roadmap."""
    doc.add_heading("7. Remediation Status & Roadmap", level=1)

    doc.add_heading("7.1 Already Mitigated", level=2)
    add_body_text(doc,
        "The following security concerns are already addressed by the platform architecture and do not require action:"
    )
    add_styled_table(doc,
        ["Concern", "Status", "Justification"],
        [
            ["Anti-CSRF Tokens (ZAP 10202)", "Mitigated", "API uses JWT bearer authentication, not cookie-based sessions. CSRF attacks require cookie-based auth to be effective."],
            ["Cookie Scoping (ZAP 90033)", "Mitigated", "Authentication cookies are managed by the Azure platform and scoped appropriately."],
            ["SQL Injection", "Mitigated", "Entity Framework Core parameterizes all queries. No raw SQL is used. SAST and DAST confirm zero injection findings."],
            ["Cross-Site Scripting", "Mitigated", "React 19 escapes all rendered content by default. No dangerouslySetInnerHTML usage detected. DAST confirms zero XSS findings."],
            ["Hardcoded Secrets", "Mitigated", "Semgrep secret detection found zero hardcoded credentials. Secrets are managed via App Service application settings."],
            ["Non-Storable Content (ZAP 10049)", "By Design", "Authenticated API responses correctly prevent caching by shared proxies."],
            ["Sec-Fetch Headers (ZAP 90005)", "Not Applicable", "These are request-side browser headers. The ZAP scanner does not send them because it is not a browser."],
            ["HSTS Header (ZAP 10035)", "Mitigated", "HSTS configured with max-age=31536000 and includeSubDomains via ASP.NET Core middleware."],
            ["Content Security Policy (ZAP 10038)", "Mitigated", "Restrictive CSP header (default-src 'none'; frame-ancestors 'none') set for all API responses."],
            ["X-Content-Type-Options (ZAP 10021)", "Mitigated", "X-Content-Type-Options: nosniff header set via security headers middleware."],
            ["X-Frame-Options (ZAP 10020)", "Mitigated", "X-Frame-Options: DENY header set via security headers middleware."],
            ["SQL Server Auditing (CKV_AZURE_23/24)", "Mitigated", "Auditing enabled with 90-day retention via Log Analytics workspace in database.bicep."],
            ["SQL Threat Detection (CKV_AZURE_25)", "Mitigated", "Advanced Threat Protection enabled with all alert types and admin email notification."],
            ["Storage Network Access (CKV_AZURE_35)", "Mitigated", "Storage Account networkAcls set to defaultAction: 'Deny' with bypass: 'AzureServices' in storage.bicep."],
            ["Managed Identity (CKV_AZURE_71)", "Mitigated", "SystemAssigned managed identity enabled on App Service with Key Vault Secrets User RBAC role."],
            ["Defender for Storage (CKV_AZURE_84)", "Mitigated", "Defender for Storage Standard tier (DefenderForStorageV2) enabled in defender.bicep."],
            ["Defender for Key Vault (CKV_AZURE_87)", "Mitigated", "Defender for Key Vault Standard tier enabled in defender.bicep."],
            ["Defender Standard Tier (CKV_AZURE_19)", "Mitigated", "Defender for SQL Standard tier enabled in defender.bicep."],
        ],
        col_widths=[2.0, 1.0, 3.8]
    )

    doc.add_heading("7.2 Accepted Risks", level=2)
    add_body_text(doc,
        "The following risks have been reviewed and accepted based on the current deployment model and cost constraints. "
        "These are genuinely low-impact items, not deferred remediation items."
    )
    add_styled_table(doc,
        ["Finding", "Rationale for Acceptance"],
        [
            ["Permissions Policy Header (ZAP 10063)", "Low risk for a backend API that does not serve HTML content or use browser features controlled by Permissions Policy."],
            ["Zone Redundancy (CKV_AZURE_225/229)", "Zone redundancy requires Premium-tier App Service Plans and database SKUs. The current B1 deployment tier prioritizes cost-effectiveness. Will be re-evaluated when scaling to a production-grade SLA deployment."],
            ["Multi-Instance Failover (CKV_AZURE_212)", "Single-instance deployment is appropriate for the current user scale. Note: B1 tier does not carry an SLA guarantee. Multi-instance with Standard tier will be configured when production traffic warrants it."],
            ["Storage Naming Convention (CKV_AZURE_43)", "The current naming convention follows the project's established pattern and meets Azure's naming requirements."],
            ["Storage Geo-Replication (CKV_AZURE_206)", "LRS (Locally Redundant Storage) is sufficient for the current deployment. GRS will be evaluated for production disaster recovery requirements."],
        ],
        col_widths=[2.5, 4.3]
    )

    doc.add_heading("7.3 Remediation Priorities", level=2)
    add_body_text(doc,
        "The following table outlines the planned remediation timeline for findings not already mitigated or accepted. "
        "This static report captures the remediation plan at a point in time. Live tracking of remediation progress "
        "is maintained in the project's GitHub Issues (label: security-report) and project board. The authoritative "
        "status of any remediation item should be verified against the issue tracker, not this document."
    )

    add_styled_table(doc,
        ["Priority", "Finding", "Target", "Owner"],
        [
            ["Medium", "Configure Entra ID SQL Auth (CKV2_AZURE_27)", "Q2 2026", "Infrastructure"],
            ["Medium", "Restrict SQL Server Public Access (CKV_AZURE_113)", "Q2 2026", "Infrastructure"],
            ["Medium", "Restrict App Service Public Access (CKV_AZURE_222)", "Q2 2026", "Infrastructure"],
            ["Medium", "Implement refresh token rotation and token revocation", "Q3 2026", "Backend"],
            ["Medium", "Migrate secrets to Azure Key Vault", "Q2 2026", "Infrastructure"],
            ["Low", "Configure Security Contact Alerts (CKV_AZURE_20-22/26-27)", "Q2 2026", "Infrastructure"],
            ["Low", "Enable HTTP/2 (CKV_AZURE_18/67)", "Q2 2026", "Infrastructure"],
            ["Low", "Configure Health Check (CKV_AZURE_213)", "Q2 2026", "Infrastructure"],
            ["Low", "Enable Entra ID Registration (CKV_AZURE_16)", "Q3 2026", "Infrastructure"],
            ["Low", "Enable Client Certificates (CKV_AZURE_17)", "Evaluate", "Infrastructure"],
        ],
        col_widths=[0.8, 3.2, 1.0, 1.5]
    )

    doc.add_heading("7.4 Recommended Additional Assessments", level=2)
    add_body_text(doc,
        "The following assessments are recommended to supplement the automated scanning pipeline and "
        "address the coverage gaps identified in Section 2.4:"
    )
    add_styled_table(doc,
        ["Assessment", "Purpose", "Recommended Frequency"],
        [
            ["Manual Penetration Test", "Identify business logic flaws, authorization bypasses, and chained vulnerabilities that automated tools miss", "Annually or before major releases"],
            ["SignalR WebSocket Security Review", "Test real-time communication layer for injection, DoS, and authorization issues", "Annually"],
            ["Horizontal Privilege Escalation Test", "Verify multi-tenant data isolation using multiple authenticated user contexts", "Annually"],
            ["Token Security Audit", "Evaluate JWT implementation, token lifetime, revocation, and refresh token handling", "Annually"],
        ],
        col_widths=[2.2, 3.0, 1.6]
    )

    doc.add_page_break()


def add_compliance_mapping(doc):
    """Section 8: Compliance Mapping."""
    doc.add_heading("8. Compliance Mapping", level=1)

    doc.add_heading("8.1 OWASP ASVS Mapping", level=2)
    add_body_text(doc,
        "The following table maps the security assessment coverage to OWASP Application Security Verification Standard (ASVS) categories."
    )

    add_styled_table(doc,
        ["ASVS Category", "Coverage", "Tools", "Gaps"],
        [
            ["V1: Architecture & Design", "Partial", "Architecture review, IaC (Checkov)", "Threat model is developer-authored (Section 2.6), not independently validated"],
            ["V2: Authentication", "Partial", "SAST (Semgrep), DAST (ZAP)", "No MFA testing, limited token lifecycle review"],
            ["V3: Session Management", "Partial", "DAST (ZAP baseline)", "No refresh token or session invalidation testing"],
            ["V4: Access Control", "Partial", "SAST (Semgrep), architecture review", "No horizontal/vertical privilege escalation testing"],
            ["V5: Validation & Encoding", "Covered", "SAST (Semgrep), DAST (ZAP active)", ""],
            ["V6: Cryptography", "Partial", "SAST (secret detection), IaC (encryption)", "No key rotation or crypto implementation review"],
            ["V7: Error Handling & Logging", "Partial", "IaC (Checkov auditing checks)", "App-level audit logging partial; SQL auditing now enabled (CKV_AZURE_23 mitigated)"],
            ["V8: Data Protection", "Partial", "SAST (Semgrep), IaC (Checkov)", "No data classification scheme defined"],
            ["V9: Communication", "Partial", "DAST (ZAP TLS checks)", "HSTS now configured (zap-10035 mitigated); SignalR WebSocket transport not tested"],
            ["V10: Malicious Code", "Covered", "SAST (Semgrep secret detection)", ""],
            ["V13: API Security", "Partial", "DAST (ZAP OpenAPI scan)", "Single user context; no authz bypass testing"],
            ["V14: Configuration", "Covered", "IaC (Checkov), DAST (ZAP headers)", ""],
        ],
        col_widths=[1.8, 0.7, 2.2, 2.1]
    )

    doc.add_heading("8.2 NIST SP 800-53 Selected Controls", level=2)
    add_body_text(doc,
        "The following NIST SP 800-53 controls are relevant to this assessment and are addressed by the automated security pipeline."
    )

    add_styled_table(doc,
        ["Control", "Title", "Assessment Coverage", "Gaps"],
        [
            ["SA-11", "Developer Testing", "SAST on every commit; DAST weekly/monthly", "No manual pen testing"],
            ["RA-5", "Vulnerability Scanning", "Automated SAST, DAST, IaC on CI/CD", "Limited C# SAST depth"],
            ["SI-10", "Input Validation", "SAST (injection), DAST (active testing)", ""],
            ["SC-8", "Transmission Confidentiality", "DAST (TLS/HSTS), IaC (HTTPS)", ""],
            ["SC-28", "Protection at Rest", "IaC (SQL TDE, Storage encryption)", "Key rotation not audited"],
            ["AC-3", "Access Enforcement", "SAST, architecture review (RBAC)", "No privilege escalation testing"],
            ["AC-4", "Information Flow", "IaC (network isolation), multi-tenancy", "Cross-tenant testing manual only"],
            ["AU-2", "Event Logging", "IaC (SQL auditing, Defender checks)", "App-level audit logging partial"],
            ["AU-6", "Audit Record Review", "GitHub Issue tracking", "No SIEM integration"],
            ["CM-6", "Configuration Settings", "IaC (Checkov CIS benchmark)", "Runtime drift not monitored"],
            ["IA-2", "User Identification", "JWT bearer token validation", "No MFA enforcement at app level"],
            ["IA-5", "Authenticator Management", "SAST (secret detection)", "No token revocation mechanism"],
        ],
        col_widths=[0.6, 1.5, 2.5, 2.2]
    )

    doc.add_heading("8.3 Continuous Monitoring", level=2)
    add_body_text(doc,
        "The Cadence security assessment is not a point-in-time audit but a continuous monitoring program integrated "
        "into the development lifecycle:"
    )

    add_styled_table(doc,
        ["Schedule", "Scan Type", "Details"],
        [
            ["Every push to main", "SAST (Semgrep)", "Full source code analysis on every code change"],
            ["Every push (infra paths)", "IaC (Checkov)", "Bicep template scanning on infrastructure changes"],
            ["Weekly (Monday 2 AM UTC)", "DAST Quick", "Authenticated baseline scan \u2014 passive analysis"],
            ["Monthly (1st, 3 AM UTC)", "DAST Full", "OpenAPI API scan \u2014 active testing of every endpoint"],
            ["Weekly (Monday 6 AM UTC)", "Consolidated Report", "Merges latest SAST + IaC + DAST results"],
            ["Monthly (1st, 7 AM UTC)", "Consolidated Report", "Captures comprehensive API scan results"],
        ],
        col_widths=[2.0, 1.5, 3.3]
    )

    add_body_text(doc,
        "Consolidated results are published as SARIF artifacts, HTML reports, and automatically opened as GitHub Issues "
        "for tracking. This ensures that new findings are visible to the development team within one week of introduction "
        "and are tracked through resolution."
    )

    doc.add_page_break()


def add_references(doc):
    """Section 9: References & Further Reading."""
    doc.add_heading("9. References & Further Reading", level=1)

    doc.add_heading("9.1 Security Testing Tools", level=2)
    add_styled_table(doc,
        ["Resource", "Description", "URL"],
        [
            ["Semgrep Documentation", "Lightweight static analysis for finding bugs and enforcing code standards", "https://semgrep.dev/docs/"],
            ["OWASP ZAP", "The world's most widely used web application security scanner", "https://www.zaproxy.org/docs/"],
            ["Checkov", "Policy-as-code tool for scanning IaC files for security misconfigurations", "https://www.checkov.io/1.Welcome/What%20is%20Checkov.html"],
        ],
        col_widths=[1.8, 3.0, 2.0]
    )

    doc.add_heading("9.2 Standards & Frameworks", level=2)
    add_styled_table(doc,
        ["Resource", "Description", "URL"],
        [
            ["OWASP Top 10 2025", "The most critical web application security risks (latest edition)", "https://owasp.org/Top10/2025/"],
            ["OWASP ASVS", "Application Security Verification Standard for secure development", "https://owasp.org/www-project-application-security-verification-standard/"],
            ["NIST SP 800-53", "Security and Privacy Controls for Information Systems", "https://csf.tools/reference/nist-sp-800-53/"],
            ["SARIF 2.1.0", "Static Analysis Results Interchange Format specification", "https://docs.oasis-open.org/sarif/sarif/v2.1.0/sarif-v2.1.0.html"],
            ["CWE", "Common Weakness Enumeration \u2014 community-developed list of software weaknesses", "https://cwe.mitre.org/"],
        ],
        col_widths=[1.8, 3.0, 2.0]
    )

    doc.add_heading("9.3 Azure Security", level=2)
    add_styled_table(doc,
        ["Resource", "Description", "URL"],
        [
            ["Azure App Service Security", "Security features and best practices for App Service", "https://learn.microsoft.com/en-us/azure/app-service/overview-security"],
            ["Azure SQL Security", "Comprehensive security overview for Azure SQL Database", "https://learn.microsoft.com/en-us/azure/azure-sql/database/security-overview"],
            ["Azure Defender for Cloud", "Cloud-native application protection platform", "https://learn.microsoft.com/en-us/azure/defender-for-cloud/"],
            ["Azure Private Endpoints", "Network isolation for Azure PaaS services", "https://learn.microsoft.com/en-us/azure/private-link/private-endpoint-overview"],
        ],
        col_widths=[1.8, 3.0, 2.0]
    )

    doc.add_page_break()


def add_appendices(doc, deduped, tools):
    """Section 10: Appendices."""
    doc.add_heading("10. Appendices", level=1)

    # Appendix A: Tool Versions
    doc.add_heading("Appendix A: Tool Versions and Configuration", level=2)
    add_styled_table(doc,
        ["Tool", "Version", "Configuration"],
        [
            ["Semgrep OSS", "1.153.1", "Rule packs: auto, csharp, typescript, react, secrets, owasp-top-ten (2025 mappings). Severity: ERROR + WARNING."],
            ["OWASP ZAP", "2.17.0", "Baseline scan (passive) + API scan (active, OpenAPI-driven). Custom rules.tsv for exclusions."],
            ["Checkov", "3.2.507", "Framework: azure. Scans all .bicep files in infrastructure/ directory."],
            ["SARIF Consolidation", "Custom", "Python script merges SAST, DAST baseline, DAST API, and IaC results into unified SARIF 2.1.0."],
        ],
        col_widths=[1.5, 1.0, 4.3]
    )

    # Appendix B: Rule Exclusions
    doc.add_heading("Appendix B: ZAP Rule Exclusions", level=2)
    add_body_text(doc,
        "The following ZAP rules are configured with custom actions in the .zap/rules.tsv file. "
        "IGNORE rules are excluded from scan results with documented rationale. WARN rules are included "
        "but at a reduced severity to ensure they are reviewed without triggering false alerts."
    )

    add_styled_table(doc,
        ["Rule ID", "Action", "Rationale"],
        [
            ["10038", "IGNORE", "Content Security Policy Header Not Set \u2014 planned for future implementation"],
            ["10063", "IGNORE", "Permissions Policy Header Not Set \u2014 low risk for backend API"],
            ["90033", "IGNORE", "Loosely Scoped Cookie \u2014 Azure platform manages auth cookie scoping"],
            ["10202", "IGNORE", "Anti-CSRF Tokens \u2014 API uses JWT bearer auth, CSRF not applicable"],
            ["10098", "WARN", "Cross-Domain Misconfiguration \u2014 review if flagged"],
            ["10020", "WARN", "X-Frame-Options Header \u2014 recommended improvement"],
            ["10021", "WARN", "X-Content-Type-Options Header Missing \u2014 recommended improvement"],
            ["10036", "WARN", "Server Leaks Version Information \u2014 cosmetic improvement"],
            ["40012", "WARN", "Cross-Site Scripting (Reflected) \u2014 review active scan findings"],
            ["40014", "WARN", "Cross-Site Scripting (Persistent) \u2014 review active scan findings"],
            ["40018", "WARN", "SQL Injection \u2014 review active scan findings"],
            ["40003", "WARN", "CRLF Injection \u2014 review active scan findings"],
        ],
        col_widths=[1.0, 1.0, 4.8]
    )

    # Appendix C: Scan Execution Details
    doc.add_heading("Appendix C: Scan Execution Details", level=2)
    add_styled_table(doc,
        ["Property", "Value"],
        [
            ["Scan Date", "March 6, 2026"],
            ["SAST Target", "Full source tree (src/)"],
            ["DAST Target", "https://app-cadence-api-uat.azurewebsites.net"],
            ["IaC Target", "infrastructure/ directory (Bicep templates)"],
            ["SARIF Schema", "https://json.schemastore.org/sarif-2.1.0.json"],
            ["Workflow", "Security Report (GitHub Actions)"],
            ["Run ID", "22767587487"],
            ["Report Format", "SARIF 2.1.0 (consolidated), HTML"],
        ],
        col_widths=[2.0, 4.8]
    )

    # Appendix D: Glossary
    doc.add_heading("Appendix D: Glossary", level=2)
    add_styled_table(doc,
        ["Term", "Definition"],
        [
            ["SAST", "Static Application Security Testing \u2014 analysis of source code without executing the program"],
            ["DAST", "Dynamic Application Security Testing \u2014 testing a running application for vulnerabilities"],
            ["IaC", "Infrastructure as Code \u2014 managing infrastructure through declarative configuration files"],
            ["SARIF", "Static Analysis Results Interchange Format \u2014 standard format for security tool outputs"],
            ["OWASP", "Open Worldwide Application Security Project \u2014 nonprofit focused on application security"],
            ["CWE", "Common Weakness Enumeration \u2014 categorized list of software and hardware weakness types"],
            ["HSEEP", "Homeland Security Exercise and Evaluation Program \u2014 doctrine for exercise design and conduct"],
            ["MSEL", "Master Scenario Events List \u2014 ordered list of scripted events (injects) for an exercise"],
            ["JWT", "JSON Web Token \u2014 compact, URL-safe means of representing claims between two parties"],
            ["HSTS", "HTTP Strict Transport Security \u2014 mechanism to enforce HTTPS connections"],
            ["TDE", "Transparent Data Encryption \u2014 Azure SQL encryption of data at rest"],
            ["CIS", "Center for Internet Security \u2014 organization providing security benchmarks and best practices"],
            ["RBAC", "Role-Based Access Control \u2014 authorization model based on user roles"],
            ["CORS", "Cross-Origin Resource Sharing \u2014 mechanism to allow restricted resources from another domain"],
            ["ASVS", "Application Security Verification Standard \u2014 OWASP framework for testing web application security controls"],
            ["FOUO", "For Official Use Only \u2014 designation for sensitive but unclassified government information"],
            ["NIST", "National Institute of Standards and Technology \u2014 U.S. agency publishing security standards and frameworks"],
            ["MFA", "Multi-Factor Authentication \u2014 requiring two or more verification factors for access"],
            ["SCA", "Software Composition Analysis \u2014 identifying known vulnerabilities in third-party dependencies"],
            ["CSP", "Content Security Policy \u2014 HTTP header that restricts sources for scripts, styles, and other resources"],
            ["TLS", "Transport Layer Security \u2014 cryptographic protocol for secure communication over a network"],
            ["CSRF", "Cross-Site Request Forgery \u2014 attack that forces authenticated users to submit unintended requests"],
            ["CDN", "Content Delivery Network \u2014 distributed network of servers for delivering web content"],
        ],
        col_widths=[1.2, 5.6]
    )

    # Appendix E: Change Log
    doc.add_heading("Appendix E: Document Change Log", level=2)
    add_body_text(doc,
        "This appendix tracks substantive changes to the report structure, methodology, and content across "
        "document versions. Minor formatting changes are not listed."
    )
    add_styled_table(doc,
        ["Version", "Date", "Changes"],
        [
            ["1.0", "2026-03-01", "Initial report generated from consolidated SARIF pipeline. "
             "Single severity model (tool-reported only). Basic architecture description. "
             "OWASP Top 10 2021 mappings. Limited compliance mapping."],
            ["1.1", "2026-03-05", "Dual severity model: assessed risk (primary) and tool severity (secondary). "
             "Vendor self-assessment disclaimer added. SAST claims tempered to acknowledge Semgrep OSS limitations. "
             "OWASP mappings updated to 2025 edition. Terminology updated (Microsoft Entra ID). "
             "Added Section 3.4 Token & Session Management. Added Section 3.5 Data Protection. "
             "Expanded Limitations & Coverage Gaps (Section 2.4). "
             "Added ASVS Gaps column (Section 8.1). Added Section 7.4 Recommended Additional Assessments. "
             "Reclassified HSTS from Low to Medium. Moved CSP from accepted risks to remediation."],
            ["1.2", "2026-03-06", "Structural improvements based on security professional review. "
             "Added Section 2.5 Architecture Review Scope (internal review disclaimer). "
             "Added Section 2.6 Threat Model (adversary profiles, impact analysis). "
             "Added Section 2.7 Data Classification (data categories and protection requirements). "
             "Section 3 cleaned to be purely descriptive\u2014gap assertions moved to Section 7 remediation. "
             "Section 7.3 updated with issue tracker reference and token lifecycle remediation items. "
             "NIST SP 800-53 table expanded to 12 controls with Gaps column (Section 8.2). "
             "Added Appendix E (this change log). Executive summary references threat model."],
            ["1.3", "2026-03-09", "Findings accuracy pass based on security professional review. "
             "CKV_AZURE_19, CKV_AZURE_84, and CKV_AZURE_87 now carry specific remediation text instead of "
             "generic boilerplate (Defender for SQL, Storage, and Key Vault are enabled in defender.bicep). "
             "ASVS gap entries updated: V1 notes threat model is developer-authored not independently validated, "
             "V7 reflects SQL auditing now enabled, V9 reflects HSTS now configured. "
             "Glossary expanded with ASVS, FOUO, NIST, MFA, SCA, CSP, TLS, CSRF, CDN."],
        ],
        col_widths=[0.6, 1.0, 5.2]
    )


# =============================================================================
# Main
# =============================================================================

def main():
    print("Parsing SARIF...")
    tools, all_findings = parse_sarif(SARIF_PATH)
    deduped = deduplicate_findings(all_findings)

    print(f"  Tools: {len(tools)}")
    print(f"  Total instances: {len(all_findings)}")
    print(f"  Unique findings: {len(deduped)}")

    print("Creating document...")
    doc = setup_document()

    print("  Cover page...")
    add_cover_page(doc)
    add_headers_footers(doc)

    print("  Table of contents...")
    add_table_of_contents(doc)

    print("  Section 1: Executive Summary...")
    add_executive_summary(doc, tools, deduped, all_findings)

    print("  Section 2: Scope & Methodology...")
    add_scope_methodology(doc)

    print("  Section 3: Security Architecture...")
    add_security_architecture(doc)

    print("  Section 4: Findings Summary...")
    add_findings_summary(doc, deduped, all_findings)

    print("  Section 5: Detailed Findings...")
    add_detailed_findings(doc, deduped)

    print("  Section 6: Risk Assessment Matrix...")
    add_risk_matrix(doc, deduped)

    print("  Section 7: Remediation Roadmap...")
    add_remediation_roadmap(doc, deduped)

    print("  Section 8: Compliance Mapping...")
    add_compliance_mapping(doc)

    print("  Section 9: References...")
    add_references(doc)

    print("  Section 10: Appendices...")
    add_appendices(doc, deduped, tools)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(OUTPUT_PATH)
    print(f"Done! File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    main()
